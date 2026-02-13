"""
Script de generation du guide developpeur au format Word (.docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1a3c6e', qn('w:val'): 'clear'})
        shading.append(bg)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r % 2 == 0:
            for c in range(len(headers)):
                shading = table.rows[r + 1].cells[c]._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f0f4f8', qn('w:val'): 'clear'})
                shading.append(bg)
    doc.add_paragraph()
    return table


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f5f5f5', qn('w:val'): 'clear'})
    shading.append(bg)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    p.add_run(text)
    return p


def add_warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'fff3cd', qn('w:val'): 'clear'})
    shading.append(bg)
    run = p.add_run('IMPORTANT : ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x7b, 0x00)
    p.add_run(text)
    return p


# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Guide Développeur')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Extranet Giffaud Groupe')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Reprise et poursuite du projet')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(6):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('12 février 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
doc.add_heading('Table des matières', level=1)

toc = [
    '1. Présentation du projet',
    '2. Ce qui a été réalisé',
    '3. Installation et lancement',
    '4. Accès distant avec ngrok',
    '5. Variables d\'environnement',
    '6. Architecture double base de données',
    '7. Structure du projet',
    '8. Applications Django',
    '9. Patterns de conception utilisés',
    '10. Commandes Django utiles',
    '11. Comment ajouter des fonctionnalités',
    '12. Points d\'attention',
    '13. Ce qui reste à faire / améliorer',
    '14. Tests',
]
for item in toc:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. PRESENTATION DU PROJET
# ============================================================
doc.add_heading('1. Présentation du projet', level=1)

add_text(
    "L'Extranet Giffaud Groupe est une application web Django permettant aux clients "
    "professionnels de consulter un catalogue de produits alimentaires et de passer "
    "des commandes en ligne. L'application s'interface avec un ERP existant via une "
    "base de données MariaDB distante en lecture seule."
)

doc.add_heading('Stack technique', level=2)
add_table(
    ['Composant', 'Technologie', 'Version'],
    [
        ['Backend', 'Python + Django', '3.x + 6.0.1'],
        ['Frontend', 'HTML, CSS, JS + Bootstrap', '5.3.2'],
        ['BDD locale', 'SQLite', '-'],
        ['BDD distante', 'MariaDB', '10.6+'],
        ['Icônes', 'Bootstrap Icons', '1.11.1'],
        ['Images', 'Pillow', '11.0.0'],
    ]
)

doc.add_heading('Dépendances Python (requirements.txt)', level=2)
add_table(
    ['Package', 'Version', 'Rôle'],
    [
        ['Django', '6.0.1', 'Framework web'],
        ['asgiref', '3.11.0', 'Interface ASGI'],
        ['sqlparse', '0.5.5', 'Parsing SQL'],
        ['tzdata', '2025.3', 'Fuseaux horaires'],
        ['Pillow', '11.0.0', 'Traitement d\'images'],
    ]
)

# ============================================================
# 2. CE QUI A ETE REALISE
# ============================================================
doc.add_page_break()
doc.add_heading('2. Ce qui a été réalisé', level=1)

doc.add_heading('Fonctionnalités côté client', level=2)
add_table(
    ['Fonctionnalité', 'État', 'Détails'],
    [
        ['Connexion / Déconnexion', 'Terminé', 'Formulaire avec toggle mot de passe, redirection selon rôle'],
        ['Réinitialisation mot de passe', 'Terminé', 'Token sécurisé par email, validité 1h, usage unique'],
        ['Profil utilisateur', 'Terminé', 'Affichage infos ERP, modification email et mot de passe'],
        ['Catalogue produits', 'Terminé', 'Grille avec prix personnalisés, pagination, recherche'],
        ['Système de filtres', 'Terminé', '30+ catégories, filtres automatiques, accordéon'],
        ['Détail produit', 'Terminé', 'Fiche complète avec catégories et description'],
        ['Produits favoris', 'Terminé', 'Top 4 produits les plus commandés'],
        ['Panier (session)', 'Terminé', 'AJAX, résumé latéral, modification quantités en temps réel'],
        ['Formulaire commande rapide', 'Terminé', 'Tableau avec tous les produits, recherche, filtres'],
        ['Validation commande', 'Terminé', 'Récapitulatif, confirmation, email, export EDI'],
        ['Historique commandes', 'Terminé', '50 dernières commandes, détails par commande'],
        ['Recommandations', 'Terminé', 'Basé sur historique d\'achats, API JSON'],
        ['Mentions légales', 'Terminé', 'Pages accessibles client et admin'],
        ['Responsive design', 'Terminé', 'Adapté mobile, tablette et desktop'],
    ]
)

doc.add_heading('Fonctionnalités côté administration', level=2)
add_table(
    ['Fonctionnalité', 'État', 'Détails'],
    [
        ['Tableau de bord', 'Terminé', 'Statistiques, activités récentes, corbeilles'],
        ['Gestion commandes', 'Terminé', 'Liste, détails, suppression douce, restauration'],
        ['Gestion utilisateurs', 'Terminé', 'CRUD complet, suppression douce, restauration'],
        ['Inscription utilisateurs', 'Terminé', 'Formulaire avec recherche client distant'],
        ['Consultation clients ERP', 'Terminé', 'Liste clients distants, cadencier prix'],
        ['Gestion demandes MDP', 'Terminé', 'Compteur global, suivi des demandes'],
        ['Profil admin', 'Terminé', 'Modification MDP, réinitialisation'],
        ['API AJAX internes', 'Terminé', 'Recherche clients, vérification MDP'],
    ]
)

doc.add_heading('Infrastructure technique', level=2)
add_table(
    ['Composant', 'État', 'Détails'],
    [
        ['Routeur double BDD', 'Terminé', 'SQLite local + MariaDB distant, routage automatique'],
        ['Export EDI/CSV', 'Terminé', 'Génération fichier ORDERS à chaque commande'],
        ['Suppression douce', 'Terminé', 'Corbeille 5 min pour commandes et utilisateurs'],
        ['Context processors', 'Terminé', 'Compteur panier + compteur demandes MDP'],
        ['Décorateur @admin_required', 'Terminé', 'Protection des vues administration'],
        ['Système de cache', 'Terminé', 'LocalMemCache, configurable Redis'],
        ['Configuration email', 'Terminé', 'SMTP en production, console en développement'],
        ['Gestion sessions', 'Terminé', 'Expiration à la fermeture du navigateur'],
    ]
)

# ============================================================
# 3. INSTALLATION ET LANCEMENT
# ============================================================
doc.add_page_break()
doc.add_heading('3. Installation et lancement', level=1)

doc.add_heading('Prérequis', level=2)
add_bullet('Python 3.10 ou supérieur')
add_bullet('pip (gestionnaire de paquets Python)')
add_bullet('Accès réseau au serveur MariaDB (pour les données produits)')
add_bullet('Git (optionnel, pour cloner le dépôt)')

doc.add_heading('Installation pas à pas', level=2)

add_text('Étape 1 : Créer l\'environnement virtuel', bold=True)
add_code('python -m venv env')

add_text('Étape 2 : Activer l\'environnement virtuel', bold=True)
add_code('env\\Scripts\\activate          # Windows')
add_code('source env/bin/activate        # Linux / Mac')

add_text('Étape 3 : Installer les dépendances', bold=True)
add_code('pip install -r requirements.txt')

add_text('Étape 4 : Configurer les variables d\'environnement', bold=True)
add_code('copy .env.example .env         # Windows')
add_code('cp .env.example .env           # Linux / Mac')
add_text('Puis éditez le fichier .env avec vos valeurs (voir section 4).')

add_text('Étape 5 : Appliquer les migrations', bold=True)
add_code('python manage.py migrate')

add_text('Étape 6 : Créer un superutilisateur (admin)', bold=True)
add_code('python manage.py createsuperuser')

add_text('Étape 7 : Lancer le serveur de développement', bold=True)
add_code('python manage.py runserver')

add_text('L\'application est accessible sur http://127.0.0.1:8000/')

add_warning(
    'Les migrations ne s\'appliquent qu\'à la base locale (SQLite). '
    'La base distante MariaDB contient des tables existantes du système ERP '
    'et ne doit pas être modifiée par Django.'
)

# ============================================================
# 4. ACCES DISTANT AVEC NGROK
# ============================================================
doc.add_page_break()
doc.add_heading('4. Accès distant avec ngrok', level=1)

add_text(
    'ngrok permet d\'exposer votre serveur de développement local sur Internet. '
    'C\'est utile pour tester depuis un mobile, montrer le site à un collègue '
    'ou tester les webhooks (emails, API externes).'
)

doc.add_heading('Installation de ngrok', level=2)

add_text('Option A — via winget (recommandé sous Windows) :', bold=True)
add_code('winget install ngrok.ngrok')

add_text('Option B — téléchargement manuel :', bold=True)
add_bullet('Aller sur https://ngrok.com/download')
add_bullet('Télécharger le fichier .zip pour Windows')
add_bullet('Extraire l\'exécutable et l\'ajouter au PATH')

doc.add_heading('Configuration (une seule fois)', level=2)

add_step(1, 'Créer un compte gratuit sur https://ngrok.com')
add_step(2, 'Récupérer votre authtoken depuis le tableau de bord ngrok.')
add_step(3, 'Enregistrer le token :')
add_code('ngrok config add-authtoken VOTRE_TOKEN_ICI')

doc.add_heading('Utilisation', level=2)

add_text(
    'Deux terminaux sont nécessaires : un pour Django, un pour ngrok.',
    bold=True
)

add_table(
    ['Terminal 1', 'Terminal 2'],
    [
        ['python manage.py runserver', 'ngrok http 8000'],
    ]
)

add_text('ngrok affiche une URL publique du type :')
add_code('Forwarding  https://xxxx-xx-xx.ngrok-free.app -> http://localhost:8000')

doc.add_heading('Configuration Django pour ngrok', level=2)

add_text(
    'Le projet est déjà configuré pour accepter les domaines ngrok. '
    'Dans extranet/settings.py, la variable CSRF_TRUSTED_ORIGINS contient :'
)
add_code(
    "CSRF_TRUSTED_ORIGINS = [\n"
    "    'https://*.ngrok-free.app',\n"
    "    'https://*.ngrok-free.dev',\n"
    "    'https://*.ngrok.io',\n"
    "]"
)

add_text(
    'Il faut également ajouter le domaine ngrok dans ALLOWED_HOSTS du fichier .env :',
)
add_code('ALLOWED_HOSTS=localhost,127.0.0.1,.ngrok-free.app,.ngrok-free.dev,.ngrok.io')

add_text(
    'Cette configuration est déjà en place dans le .env du projet. '
    'Aucune modification n\'est nécessaire.'
)

doc.add_heading('Vérification', level=2)
add_step(1, 'Lancez le serveur Django dans le premier terminal.')
add_step(2, 'Lancez ngrok http 8000 dans le second terminal.')
add_step(3, 'Copiez l\'URL https://...ngrok-free.app affichée par ngrok.')
add_step(4, 'Ouvrez cette URL dans un navigateur ou sur un téléphone.')
add_step(5, 'Vous devriez voir la page de connexion de l\'Extranet.')

add_warning(
    'L\'URL ngrok change à chaque redémarrage (version gratuite). '
    'Pour un domaine fixe, un abonnement ngrok payant est nécessaire.'
)

# ============================================================
# 5. VARIABLES D'ENVIRONNEMENT
# ============================================================
doc.add_page_break()
doc.add_heading('5. Variables d\'environnement', level=1)

add_text(
    'Le fichier .env à la racine du projet contient toutes les configurations sensibles. '
    'Il ne doit jamais être versionné dans Git.'
)

doc.add_heading('Configuration Django', level=2)
add_table(
    ['Variable', 'Description', 'Valeur dev', 'Valeur prod'],
    [
        ['SECRET_KEY', 'Clé secrète (obligatoire)', 'django-insecure-xxx', 'Clé aléatoire longue'],
        ['DEBUG', 'Mode debug', 'True', 'False'],
        ['ALLOWED_HOSTS', 'Hôtes autorisés', 'localhost,127.0.0.1', 'votre-domaine.com'],
    ]
)

doc.add_heading('Base de données distante (MariaDB)', level=2)
add_table(
    ['Variable', 'Description', 'Exemple'],
    [
        ['DB_LOGIGVD_NAME', 'Nom de la base', 'logigvd'],
        ['DB_LOGIGVD_USER', 'Utilisateur', 'extranet_user'],
        ['DB_LOGIGVD_PASSWORD', 'Mot de passe', '********'],
        ['DB_LOGIGVD_HOST', 'Adresse du serveur', '192.168.1.100'],
        ['DB_LOGIGVD_PORT', 'Port', '3306'],
    ]
)


# ============================================================
# 6. ARCHITECTURE DOUBLE BDD
# ============================================================
doc.add_page_break()
doc.add_heading('6. Architecture double base de données', level=1)

doc.add_heading('Principe', level=2)
add_text(
    'Le projet utilise deux bases de données simultanément, gérées par un routeur '
    'Django personnalisé (extranet/db_router.py).'
)

add_table(
    ['Base', 'Moteur', 'Usage', 'Accès'],
    [
        ['default', 'SQLite', 'Utilisateurs, commandes, recommandations, historiques', 'Lecture / Écriture'],
        ['logigvd', 'MariaDB', 'Produits, clients, prix, catalogue', 'Lecture seule'],
    ]
)

doc.add_heading('Routeur de base de données', level=2)
add_text('Le fichier extranet/db_router.py contient la classe DatabaseRouter qui :')
add_bullet('dirige les lectures/écritures des modèles vers la bonne base,')
add_bullet('empêche toute écriture sur la base distante (logigvd),')
add_bullet('empêche les migrations sur la base distante.')

add_text('Tables routées vers MariaDB :', bold=True)
add_code("LOGIGVD_TABLES = {'comcli', 'comclilig', 'catalogue', 'prod'}")

doc.add_heading('Modèles distants', level=2)
add_text(
    'Les modèles connectés à la base distante sont définis dans catalogue/models.py. '
    'Ils ont tous managed = False pour empêcher Django de créer ou modifier leurs tables.'
)
add_table(
    ['Modèle', 'Table', 'Description'],
    [
        ['Prod', 'prod', 'Catalogue des produits (code, libellé)'],
        ['ComCli', 'comcli', 'Clients et adresses de livraison'],
        ['ComCliLig', 'comclilig', 'Historique des lignes de commande avec prix'],
        ['Catalogue', 'catalogue', 'Disponibilité produit/client (clé composite)'],
    ]
)

doc.add_heading('Requêtes sur la base distante', level=2)
add_text(
    'Les requêtes vers la base distante sont centralisées dans catalogue/services.py. '
    'La fonction principale get_produits_client(utilisateur) exécute des requêtes SQL '
    'optimisées avec JOINs pour récupérer les produits disponibles avec leurs prix.'
)
add_warning(
    'Ne pas utiliser .save(), .delete() ou .create() sur les modèles distants. '
    'Toute tentative d\'écriture sera bloquée par le routeur.'
)

# ============================================================
# 7. STRUCTURE DU PROJET
# ============================================================
doc.add_page_break()
doc.add_heading('7. Structure du projet', level=1)

add_code(
    "Code/\n"
    "├── manage.py                    # Point d'entrée Django\n"
    "├── requirements.txt             # Dépendances Python\n"
    "├── .env / .env.example          # Variables d'environnement\n"
    "├── db.sqlite3                   # Base locale\n"
    "│\n"
    "├── extranet/                    # Configuration Django\n"
    "│   ├── settings.py              # Configuration principale\n"
    "│   ├── urls.py                  # Routage racine\n"
    "│   ├── wsgi.py / asgi.py        # Points d'entrée serveur\n"
    "│   └── db_router.py             # Routeur multi-bases\n"
    "│\n"
    "├── clients/                     # Authentification & profils\n"
    "│   ├── models.py, views.py, forms.py, urls.py\n"
    "│   └── context_processors.py    # Compteur demandes MDP\n"
    "│\n"
    "├── catalogue/                   # Catalogue produits\n"
    "│   ├── models.py               # Modèles distants (managed=False)\n"
    "│   ├── services.py             # Requêtes BDD, filtres\n"
    "│   └── views.py, urls.py\n"
    "│\n"
    "├── commandes/                   # Panier & commandes\n"
    "│   ├── models.py               # Commande, LigneCommande\n"
    "│   ├── services.py             # Export EDI/CSV\n"
    "│   ├── context_processors.py   # Compteur panier\n"
    "│   └── views.py, urls.py\n"
    "│\n"
    "├── recommandations/             # Recommandations\n"
    "│   ├── models.py               # HistoriqueAchat, Préférences\n"
    "│   └── services.py, views.py, urls.py\n"
    "│\n"
    "├── administration/              # Back-office admin\n"
    "│   ├── urls.py\n"
    "│   └── views/                   # Vues par module\n"
    "│       ├── dashboard.py\n"
    "│       ├── utils/decorators.py  # @admin_required\n"
    "│       ├── utils/filtres.py     # Filtres admin\n"
    "│       ├── commandes/           # Gestion commandes\n"
    "│       ├── utilisateurs/        # Gestion utilisateurs\n"
    "│       ├── clients/             # Clients distants\n"
    "│       └── api/                 # Endpoints AJAX\n"
    "│\n"
    "├── templates/\n"
    "│   ├── cote_client/             # Templates client\n"
    "│   └── administration/          # Templates admin\n"
    "│\n"
    "├── static/\n"
    "│   ├── css/style.css, style_admin.css\n"
    "│   └── js/ (panier.js, commander.js, ...)\n"
    "│\n"
    "└── edi_exports/                 # Fichiers EDI générés"
)

# ============================================================
# 8. APPLICATIONS DJANGO
# ============================================================
doc.add_page_break()
doc.add_heading('8. Applications Django', level=1)

doc.add_heading('clients', level=2)
add_text('Gère l\'authentification, les profils et la réinitialisation de mot de passe.', bold=True)
add_table(
    ['Fichier', 'Contenu clé'],
    [
        ['models.py', 'Utilisateur (profil), TokenResetPassword, DemandeMotDePasse, UtilisateurSupprime'],
        ['views.py', 'connexion, deconnexion, profil, modifier_mot_de_passe, modifier_email, reset MDP'],
        ['forms.py', 'ConnexionForm, formulaires de modification'],
        ['context_processors.py', 'Compteur de demandes MDP en attente (nb_demandes_mdp_global)'],
    ]
)
add_text(
    'Le modèle Utilisateur lie le User Django à un code_tiers (code client ERP). '
    'La méthode get_client_distant() récupère les données client depuis MariaDB.'
)

doc.add_heading('catalogue', level=2)
add_text('Consultation du catalogue de produits depuis la base distante.', bold=True)
add_table(
    ['Fichier', 'Contenu clé'],
    [
        ['models.py', 'Prod, ComCli, ComCliLig, Catalogue (tous managed=False)'],
        ['services.py', 'get_produits_client(), get_produit_by_reference(), système de filtres'],
        ['views.py', 'liste_produits, detail_produit, favoris, commander, mentions_legales'],
    ]
)
add_text(
    'Le fichier services.py contient FILTRES_DISPONIBLES, un dictionnaire de 30+ catégories '
    'de filtres (viandes, découpes, charcuterie, etc.) et la logique de filtrage automatique.'
)

doc.add_heading('commandes', level=2)
add_text('Gestion du panier et des commandes.', bold=True)
add_table(
    ['Fichier', 'Contenu clé'],
    [
        ['models.py', 'Commande, LigneCommande, CommandeSupprimee, HistoriqueSuppression'],
        ['services.py', 'generer_csv_edi() — génération du fichier EDI, envoyer_commande() — stub ERP'],
        ['views.py', 'voir_panier, ajouter/modifier/supprimer, valider_commande, historique'],
        ['context_processors.py', 'Compteur panier (panier_count)'],
    ]
)
add_text(
    'Le panier est stocké en session (request.session[\'panier\']). '
    'À la validation, les prix sont copiés dans LigneCommande (pattern snapshot).'
)

doc.add_heading('recommandations', level=2)
add_text('Système de recommandations basé sur l\'historique d\'achats.', bold=True)
add_table(
    ['Fichier', 'Contenu clé'],
    [
        ['models.py', 'HistoriqueAchat (quantité cumulée, fréquence), PreferenceCategorie (scores)'],
        ['services.py', 'Algorithme de recommandation'],
        ['views.py', 'Page recommandations + API JSON (/api/ et /api/favoris/)'],
    ]
)

doc.add_heading('administration', level=2)
add_text('Interface d\'administration personnalisée (pas le Django admin).', bold=True)
add_text(
    'Les vues sont organisées en sous-dossiers dans administration/views/ et centralisées '
    'via administration/views/__init__.py.'
)
add_table(
    ['Dossier', 'Contenu'],
    [
        ['views/dashboard.py', 'Tableau de bord avec statistiques'],
        ['views/commandes/', 'CRUD commandes (liste, détails, supprimer, restaurer)'],
        ['views/utilisateurs/', 'CRUD utilisateurs (liste, inscription, modifier, supprimer)'],
        ['views/clients/', 'Consultation clients distants, cadencier prix'],
        ['views/api/', 'Endpoints AJAX (recherche clients, vérification MDP)'],
        ['views/auth/', 'Profil admin, changement MDP'],
        ['views/utils/decorators.py', 'Décorateur @admin_required (vérifie is_staff)'],
        ['views/utils/filtres.py', 'Système de filtres côté admin'],
    ]
)

# ============================================================
# 9. PATTERNS DE CONCEPTION
# ============================================================
doc.add_page_break()
doc.add_heading('9. Patterns de conception utilisés', level=1)

doc.add_heading('Routage multi-bases', level=2)
add_text(
    'Le DatabaseRouter (extranet/db_router.py) redirige automatiquement chaque requête ORM '
    'vers la bonne base selon le modèle. Pas besoin d\'appeler .using() manuellement '
    'dans la plupart des cas.'
)

doc.add_heading('Suppression douce (Soft Delete)', level=2)
add_text(
    'Les commandes et utilisateurs ne sont pas supprimés immédiatement. Ils sont déplacés '
    'dans une table "corbeille" pendant 5 minutes avant suppression définitive.'
)
add_code(
    "Élément actif\n"
    "  → Suppression douce → Corbeille (5 min)\n"
    "     → Restauration possible\n"
    "     → Après 5 min : suppression définitive + archivage"
)
add_text('Tables impliquées :', bold=True)
add_bullet('Commande → CommandeSupprimee → HistoriqueSuppression')
add_bullet('Utilisateur → UtilisateurSupprime → HistoriqueSuppressionUtilisateur')

doc.add_heading('Panier en session', level=2)
add_text(
    'Le panier est stocké dans la session Django (dictionnaire Python). '
    'Avantages : pas de table BDD, performances optimales, vidé automatiquement à la fermeture.'
)
add_code(
    "session['panier'] = {\n"
    "    'PROD001': {'reference': 'PROD001', 'nom': '...', 'quantite': 3, 'prix_unitaire': 12.50}\n"
    "}"
)

doc.add_heading('Pattern Snapshot (lignes de commande)', level=2)
add_text(
    'Quand une commande est validée, le nom et le prix de chaque produit sont copiés '
    'dans LigneCommande. Ainsi, même si le catalogue évolue, l\'historique reste fidèle.'
)

doc.add_heading('Couche de services', level=2)
add_text(
    'La logique métier est séparée des vues dans des fichiers services.py :'
)
add_bullet('requêtes BDD distante et système de filtres', bold_prefix='catalogue/services.py : ')
add_bullet('génération EDI/CSV et transmission commande', bold_prefix='commandes/services.py : ')
add_bullet('algorithme de recommandation', bold_prefix='recommandations/services.py : ')

doc.add_heading('Context processors', level=2)
add_text(
    'Deux context processors injectent des variables dans tous les templates :'
)
add_bullet('nombre d\'articles dans le panier', bold_prefix='panier_count : ')
add_bullet('nombre de demandes MDP en attente (admin)', bold_prefix='nb_demandes_mdp_global : ')

# ============================================================
# 10. COMMANDES DJANGO UTILES
# ============================================================
doc.add_page_break()
doc.add_heading('10. Commandes Django utiles', level=1)

add_table(
    ['Commande', 'Description'],
    [
        ['python manage.py runserver', 'Lance le serveur de développement'],
        ['python manage.py migrate', 'Applique les migrations (BDD locale)'],
        ['python manage.py makemigrations [app]', 'Crée les fichiers de migration'],
        ['python manage.py createsuperuser', 'Crée un superutilisateur'],
        ['python manage.py shell', 'Console Python avec contexte Django'],
        ['python manage.py dbshell', 'Console SQL de la base locale'],
        ['python manage.py collectstatic', 'Collecte les fichiers statiques (production)'],
        ['python manage.py showmigrations', 'Affiche l\'état des migrations'],
        ['python manage.py check', 'Vérifie la configuration du projet'],
        ['python manage.py test', 'Lance les tests'],
        ['python manage.py test [app]', 'Lance les tests d\'une application'],
    ]
)

# ============================================================
# 11. COMMENT AJOUTER DES FONCTIONNALITES
# ============================================================
doc.add_page_break()
doc.add_heading('11. Comment ajouter des fonctionnalités', level=1)

doc.add_heading('Ajouter un filtre de produits', level=2)
add_step(1, 'Ouvrir catalogue/services.py.')
add_step(2, 'Trouver le dictionnaire FILTRES_DISPONIBLES.')
add_step(3, 'Ajouter une entrée dans le groupe approprié :')
add_code(
    '"code_filtre": {\n'
    '    "label": "Nom affiché",\n'
    '    "termes": ["mot1", "mot2"]  # Mots-clés cherchés dans le libellé\n'
    '}'
)
add_step(4, 'Les filtres apparaîtront automatiquement dans le catalogue.')

doc.add_heading('Ajouter une vue administration', level=2)
add_step(1, 'Créer le fichier dans administration/views/[module]/.')
add_step(2, 'Décorer la vue avec @admin_required.')
add_step(3, 'Importer la vue dans administration/views/__init__.py.')
add_step(4, 'Ajouter la route dans administration/urls.py.')
add_step(5, 'Créer le template dans templates/administration/[module]/.')

doc.add_heading('Modifier le processus de commande', level=2)
add_step(1, 'Modifier la vue dans commandes/views.py (fonction valider_commande).')
add_step(2, 'Si changement de modèle, modifier commandes/models.py.')
add_step(3, 'Créer les migrations : python manage.py makemigrations commandes.')
add_step(4, 'Appliquer : python manage.py migrate.')
add_step(5, 'Adapter le format EDI dans commandes/services.py si nécessaire.')

doc.add_heading('Connecter l\'ERP pour l\'envoi de commandes', level=2)
add_text(
    'La fonction envoyer_commande() dans commandes/services.py est actuellement un stub. '
    'C\'est le point d\'intégration principal avec votre ERP.'
)
add_code(
    "def envoyer_commande(commande_data):\n"
    "    # TODO: Implémenter selon votre logiciel\n"
    "    # Options :\n"
    "    # - API REST : requests.post(url, json=commande_data)\n"
    "    # - SOAP/XML : zeep ou suds\n"
    "    # - Base de données : insert dans une table ERP\n"
    "    # - Email : envoyer le fichier EDI par email\n"
    "    # - SFTP : uploader le fichier EDI\n"
    "    pass"
)

# ============================================================
# 12. POINTS D'ATTENTION
# ============================================================
doc.add_page_break()
doc.add_heading('12. Points d\'attention', level=1)

doc.add_heading('Base distante en lecture seule', level=2)
add_text(
    'Ne jamais tenter d\'écrire dans la base MariaDB (logigvd). Le routeur bloque les écritures, '
    'mais il vaut mieux ne pas essayer. Les modèles distants ont managed=False.'
)

doc.add_heading('Migrations', level=2)
add_text(
    'Les commandes makemigrations et migrate ne concernent que la base locale (SQLite). '
    'Après modification d\'un modèle local, toujours créer et appliquer la migration.'
)

doc.add_heading('Fichier .env', level=2)
add_text(
    'Le fichier .env contient des secrets (clé Django, mots de passe BDD). '
    'Il est dans le .gitignore et ne doit jamais être commité.'
)

doc.add_heading('Export EDI', level=2)
add_text(
    'Les fichiers EDI sont générés dans le dossier edi_exports/ à la racine du dépôt parent. '
    'En production, vérifier que ce dossier existe et que l\'application a les droits d\'écriture.'
)
add_code("Chemin : C:\\Users\\Giffaud\\Documents\\Site_extranet\\edi_exports\\")

doc.add_heading('Nettoyage de la corbeille', level=2)
add_text(
    'Les méthodes nettoyer_anciens() / nettoyer_anciennes() suppriment les éléments expirés '
    'de la corbeille. Elles sont appelées lors de l\'affichage du dashboard admin, mais en '
    'production il serait préférable de les planifier via une tâche cron ou Celery.'
)

doc.add_heading('Sessions et panier', level=2)
add_text(
    'Le panier est en session. SESSION_EXPIRE_AT_BROWSER_CLOSE = True signifie que '
    'le panier est perdu si l\'utilisateur ferme son navigateur. C\'est un comportement voulu.'
)

doc.add_heading('Tests', level=2)
add_text(
    'Les fichiers tests.py existent dans chaque application mais sont vides. '
    'Il est recommandé d\'écrire des tests avant toute nouvelle fonctionnalité.'
)

# ============================================================
# 13. CE QUI RESTE A FAIRE / AMELIORER
# ============================================================
doc.add_page_break()
doc.add_heading('13. Ce qui reste à faire / améliorer', level=1)

doc.add_heading('Intégrations à compléter', level=2)
add_table(
    ['Tâche', 'Priorité', 'Détails'],
    [
        ['Implémenter envoyer_commande()', 'Haute', 'Stub dans commandes/services.py — connecter à l\'ERP'],
        ['Tâche planifiée de nettoyage', 'Moyenne', 'Automatiser nettoyer_anciens() via cron ou Celery'],
        ['Tests unitaires', 'Fait', '122 tests couvrant modèles, vues, formulaires et utilitaires (voir section 14)'],
    ]
)

doc.add_heading('Améliorations possibles', level=2)
add_table(
    ['Amélioration', 'Description'],
    [
        ['Cache Redis', 'Remplacer LocalMemCache par Redis en production (settings.py prêt)'],
        ['Notifications email avancées', 'Envoyer des rappels, confirmations d\'expédition, etc.'],
        ['Export PDF des commandes', 'Générer des bons de commande au format PDF'],
        ['Tableau de bord analytics', 'Statistiques avancées (graphiques, tendances)'],
        ['API REST complète', 'Exposer une API REST avec Django REST Framework'],
        ['Recherche avancée', 'Implémenter une recherche full-text (Elasticsearch, PostgreSQL)'],
        ['Gestion des stocks', 'Afficher les disponibilités temps réel depuis l\'ERP'],
        ['Application mobile', 'PWA ou application native connectée à l\'API'],
        ['Logs et monitoring', 'Intégrer Sentry ou un système de logs centralisé'],
        ['Déploiement CI/CD', 'Pipeline GitHub Actions ou GitLab CI'],
    ]
)

doc.add_heading('Configuration production', level=2)
add_text('Checklist avant mise en production :', bold=True)
add_bullet('DEBUG = False')
add_bullet('SECRET_KEY aléatoire et sécurisée')
add_bullet('ALLOWED_HOSTS avec les domaines de production')
add_bullet('CSRF_TRUSTED_ORIGINS configuré')
add_bullet('Configurer STATIC_ROOT + collectstatic')
add_bullet('Serveur WSGI (Gunicorn) derrière un reverse proxy (Nginx)')
add_bullet('Certificat SSL / HTTPS')
add_bullet('Sauvegardes automatiques de db.sqlite3')
add_bullet('Dossier edi_exports/ avec permissions adéquates')
add_bullet('Serveur SMTP de production configuré')

# ============================================================
# 14. TESTS
# ============================================================
doc.add_page_break()
doc.add_heading('14. Tests', level=1)

add_text(
    'Le projet dispose de 122 tests unitaires couvrant l\'ensemble des applications. '
    'Les tests vérifient les modèles, les vues, les formulaires, les décorateurs, '
    'le routeur de base de données et les utilitaires.'
)

doc.add_heading('Lancer les tests', level=2)
add_text('Pour exécuter tous les tests :')
add_code('python manage.py test')
add_text('Pour tester une application spécifique :')
add_code('python manage.py test clients')
add_code('python manage.py test commandes')
add_code('python manage.py test catalogue')
add_code('python manage.py test recommandations')
add_code('python manage.py test administration')

doc.add_heading('Couverture par application', level=2)
add_table(
    ['Application', 'Nb tests', 'Éléments testés'],
    [
        ['clients', '37', 'Modèles (Utilisateur, TokenResetPassword, UtilisateurSupprime, '
         'HistoriqueSuppressionUtilisateur), formulaire ConnexionForm, vues (connexion, '
         'déconnexion, contact, modifier mot de passe, modifier email, reset password)'],
        ['commandes', '30', 'Modèles (Commande, LigneCommande, CommandeSupprimee, '
         'HistoriqueSuppression), utilitaires (parse_date), session panier, '
         'context processor, vues (panier, historique, détails, confirmation)'],
        ['catalogue', '4', 'Vues mentions légales, contrôle d\'accès (redirections '
         'pour utilisateurs non connectés). Les modèles managed=False ne sont pas '
         'testables directement (base distante MariaDB)'],
        ['recommandations', '14', 'Modèles (HistoriqueAchat avec enregistrer_achat, '
         'PreferenceCategorie avec unicité et ordering), contrôle d\'accès aux vues'],
        ['administration', '24', 'Décorateur is_admin, routeur DatabaseRouter '
         '(lecture/écriture/migration), contrôle d\'accès (non connecté, client, admin), '
         'accès aux pages admin (commandes, utilisateurs, inscription, profil, '
         'reset password, mentions légales)'],
    ]
)

doc.add_heading('Stratégie de mock pour logigvd', level=2)
add_text(
    'Certaines vues appellent get_client_distant() qui interroge la base MariaDB distante '
    '(logigvd). En environnement de test, cette base n\'est pas disponible. '
    'Les tests utilisent unittest.mock.patch pour simuler ces appels :'
)
add_code(
    'from types import SimpleNamespace\n'
    'from unittest.mock import patch\n'
    '\n'
    'MOCK_CLIENT_DISTANT = SimpleNamespace(\n'
    '    nom=\'CLIENT TEST\', complement=\'\',\n'
    '    adresse=\'1 rue Test\', cp=\'44000\',\n'
    '    acheminement=\'NANTES\'\n'
    ')\n'
    '\n'
    '@patch(\'clients.views.get_client_distant\',\n'
    '       return_value=MOCK_CLIENT_DISTANT)\n'
    'class MaVueTest(TestCase):\n'
    '    ...'
)
add_text(
    'Important : le patch doit cibler le module où la fonction est importée '
    '(ex: clients.views.get_client_distant), et non le module d\'origine '
    '(catalogue.services.get_client_distant).',
    bold=True
)

doc.add_heading('Conventions de test', level=2)
add_bullet('Chaque fichier tests.py contient des fonctions helpers (creer_utilisateur, creer_admin)')
add_bullet('Les tests de vues vérifient les codes HTTP (200, 302) et les redirections')
add_bullet('Les tests de modèles vérifient la création, les méthodes, le __str__, et les contraintes d\'unicité')
add_bullet('Les vues accédant à logigvd sont systématiquement mockées avec @patch')
add_bullet('Le dashboard admin (SQL brut vers logigvd) n\'est pas testé unitairement')

# ============================================================
# PIED DE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Document généré le 12/02/2026 — Extranet Giffaud Groupe')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ============================================================
# SAUVEGARDE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'GUIDE_DEVELOPPEUR.docx')
doc.save(output_path)
print(f"Document généré : {output_path}")
