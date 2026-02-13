"""
Script de generation de la documentation technique au format Word (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_table(headers, rows):
    """Ajoute un tableau formate."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-tete
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1a3c6e',
            qn('w:val'): 'clear'
        })
        shading.append(bg)

    # Lignes
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        # Alternance couleurs
        if r % 2 == 0:
            for c in range(len(headers)):
                shading = table.rows[r + 1].cells[c]._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'f0f4f8',
                    qn('w:val'): 'clear'
                })
                shading.append(bg)

    doc.add_paragraph()
    return table


def add_code_block(text):
    """Ajoute un bloc de code formate."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    # Fond gris
    shading = p._element.get_or_add_pPr()
    bg = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'f5f5f5',
        qn('w:val'): 'clear'
    })
    shading.append(bg)
    return p


def add_bullet(text, bold_prefix=None):
    """Ajoute un point avec optionnel texte en gras."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_text(text, bold=False):
    """Ajoute un paragraphe de texte."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Documentation Technique')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Extranet Giffaud Groupe')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Plateforme de commande en ligne B2B')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(6):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('12 février 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
doc.add_heading('Table des matières', level=1)

toc_items = [
    '1. Présentation générale',
    '2. Architecture du projet',
    '3. Structure des répertoires',
    '4. Configuration',
    '5. Base de données',
    '6. Modèles de données',
    '7. Routage URL',
    '8. Vues et logique métier',
    '9. Templates',
    '10. Fichiers statiques',
    '11. Authentification et autorisation',
    '12. Points d\'entrée API',
    '13. Système de recommandations',
    '14. Gestion du panier',
    '15. Export EDI',
    '16. Suppression douce et restauration',
    '17. Système de filtres',
    '18. Dépendances',
    '19. Variables d\'environnement',
    '20. Déploiement',
]

for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. PRESENTATION GENERALE
# ============================================================
doc.add_heading('1. Présentation générale', level=1)

doc.add_heading('Description', level=2)
add_text(
    "L'Extranet Giffaud Groupe est une plateforme de commande en ligne B2B (Business-to-Business) "
    "destinée aux clients professionnels du groupe Giffaud. Elle permet de consulter un catalogue "
    "de produits alimentaires (viandes, charcuterie, fromages, etc.), de passer des commandes "
    "et de suivre l'historique des achats."
)

doc.add_heading('Stack technique', level=2)
add_table(
    ['Composant', 'Technologie'],
    [
        ['Backend', 'Python 3 + Django 6.0.1'],
        ['Frontend', 'HTML5, CSS3, JavaScript'],
        ['Framework CSS', 'Bootstrap 5.3.2'],
        ['Icônes', 'Bootstrap Icons 1.11.1'],
        ['BDD locale', 'SQLite'],
        ['BDD distante', 'MariaDB'],
        ['Serveur WSGI', 'Compatible Gunicorn / mod_wsgi'],
    ]
)

# ============================================================
# 2. ARCHITECTURE DU PROJET
# ============================================================
doc.add_page_break()
doc.add_heading('2. Architecture du projet', level=1)

doc.add_heading('Schéma global', level=2)
add_code_block(
    "Navigateur Client\n"
    "       |\n"
    "       v\n"
    "  [Django Application]\n"
    "       |\n"
    "       +---> [SQLite - BDD locale]\n"
    "       |       Utilisateurs, Commandes,\n"
    "       |       Recommandations, Historiques\n"
    "       |\n"
    "       +---> [MariaDB - BDD distante (lecture seule)]\n"
    "               Produits, Clients, Lignes historiques, Catalogue"
)

doc.add_heading('Applications Django', level=2)
add_text("Le projet est composé de 5 applications :")
add_table(
    ['Application', 'Rôle'],
    [
        ['clients', 'Authentification, profils utilisateurs, tokens de réinitialisation'],
        ['catalogue', 'Consultation des produits (BDD distante)'],
        ['commandes', 'Panier, validation, historique des commandes'],
        ['recommandations', 'Suggestions personnalisées basées sur l\'historique d\'achats'],
        ['administration', 'Tableau de bord, gestion utilisateurs et commandes'],
    ]
)

doc.add_heading('Routeur de base de données', level=2)
add_text(
    "Le fichier extranet/db_router.py dirige automatiquement les requêtes vers la bonne base de données :"
)
add_bullet("Utilisateur, Commande, LigneCommande, HistoriqueAchat, etc.", bold_prefix="SQLite (default) : ")
add_bullet("Prod, ComCli, ComCliLig, Catalogue — tous en managed=False", bold_prefix="MariaDB (logigvd) : ")

# ============================================================
# 3. STRUCTURE DES REPERTOIRES
# ============================================================
doc.add_page_break()
doc.add_heading('3. Structure des répertoires', level=1)

add_code_block(
    "Code/\n"
    "├── manage.py                    # Point d'entrée Django\n"
    "├── requirements.txt             # Dépendances Python\n"
    "├── .env                         # Variables d'environnement\n"
    "├── db.sqlite3                   # Base SQLite locale\n"
    "│\n"
    "├── extranet/                    # Configuration Django\n"
    "│   ├── settings.py              # Configuration principale\n"
    "│   ├── urls.py                  # Routage racine\n"
    "│   ├── wsgi.py / asgi.py        # Points d'entrée serveur\n"
    "│   └── db_router.py             # Routeur multi-bases\n"
    "│\n"
    "├── clients/                     # App : Authentification & profils\n"
    "├── catalogue/                   # App : Catalogue produits\n"
    "├── commandes/                   # App : Panier & commandes\n"
    "├── recommandations/             # App : Recommandations\n"
    "├── administration/              # App : Back-office admin\n"
    "│   └── views/                   # Vues organisées par module\n"
    "│       ├── dashboard.py\n"
    "│       ├── utils/               # Décorateurs et filtres\n"
    "│       ├── commandes/           # Gestion commandes\n"
    "│       ├── utilisateurs/        # Gestion utilisateurs\n"
    "│       ├── clients/             # Consultation clients distants\n"
    "│       ├── api/                 # Endpoints AJAX\n"
    "│       ├── auth/                # Profil et MDP admin\n"
    "│       └── legal/               # Mentions légales\n"
    "│\n"
    "├── templates/\n"
    "│   ├── cote_client/             # Templates front client\n"
    "│   └── administration/          # Templates back-office\n"
    "│\n"
    "├── static/\n"
    "│   ├── css/                     # Feuilles de style\n"
    "│   ├── images/                  # Images\n"
    "│   └── js/                      # Scripts JavaScript\n"
    "│\n"
    "├── edi_exports/                 # Fichiers EDI générés\n"
    "└── media/                       # Fichiers uploadés"
)

# ============================================================
# 4. CONFIGURATION
# ============================================================
doc.add_page_break()
doc.add_heading('4. Configuration', level=1)

doc.add_heading('Applications installées', level=2)
add_code_block(
    "INSTALLED_APPS = [\n"
    "    'django.contrib.admin',\n"
    "    'django.contrib.auth',\n"
    "    'django.contrib.contenttypes',\n"
    "    'django.contrib.sessions',\n"
    "    'django.contrib.messages',\n"
    "    'django.contrib.staticfiles',\n"
    "    'clients',\n"
    "    'catalogue',\n"
    "    'commandes',\n"
    "    'recommandations',\n"
    "    'administration',\n"
    "]"
)

doc.add_heading('Middleware', level=2)
add_table(
    ['Middleware', 'Rôle'],
    [
        ['SecurityMiddleware', 'En-têtes de sécurité HTTP'],
        ['SessionMiddleware', 'Gestion des sessions'],
        ['CommonMiddleware', 'Utilitaires communs'],
        ['CsrfViewMiddleware', 'Protection CSRF'],
        ['AuthenticationMiddleware', 'Authentification utilisateur'],
        ['MessageMiddleware', 'Messages flash'],
        ['XFrameOptionsMiddleware', 'Protection clickjacking'],
    ]
)

doc.add_heading('Redirections d\'authentification', level=2)
add_table(
    ['Paramètre', 'Valeur'],
    [
        ['LOGIN_URL', 'clients:connexion'],
        ['LOGIN_REDIRECT_URL', 'catalogue:liste'],
        ['LOGOUT_REDIRECT_URL', 'clients:connexion'],
    ]
)

doc.add_heading('Cache', level=2)
add_bullet("5 minutes de durée de vie", bold_prefix="LocalMemCache : ")

doc.add_heading('Internationalisation', level=2)
add_bullet("fr-fr", bold_prefix="Langue : ")
add_bullet("Europe/Paris", bold_prefix="Fuseau horaire : ")

doc.add_heading('Session', level=2)
add_text("SESSION_EXPIRE_AT_BROWSER_CLOSE = True — La session expire à la fermeture du navigateur.")

# ============================================================
# 5. BASE DE DONNEES
# ============================================================
doc.add_page_break()
doc.add_heading('5. Base de données', level=1)

doc.add_heading('Architecture multi-bases', level=2)
add_text("Le projet utilise deux bases de données simultanément :")

doc.add_heading('Base locale — SQLite (default)', level=3)
add_text("Stocke toutes les données métier de l'application :")
add_bullet("Utilisateurs et profils")
add_bullet("Commandes et lignes de commandes")
add_bullet("Historiques d'achats et préférences")
add_bullet("Corbeilles (suppression douce)")
add_bullet("Tokens de réinitialisation de mot de passe")

doc.add_heading('Base distante — MariaDB (logigvd)', level=3)
add_text("Base en lecture seule connectée au système ERP existant :")
add_table(
    ['Table', 'Description'],
    [
        ['prod', 'Catalogue des produits'],
        ['comcli', 'Clients et adresses de livraison'],
        ['comclilig', 'Historique des lignes de commande (prix)'],
        ['catalogue', 'Disponibilité produit/client'],
    ]
)

doc.add_heading('Routeur (extranet/db_router.py)', level=3)
add_text(
    "Le routeur dirige automatiquement les requêtes vers la bonne base. "
    "Il empêche toute écriture sur la base distante (allow_migrate retourne False pour logigvd)."
)
add_code_block("LOGIGVD_TABLES = {'comcli', 'comclilig', 'catalogue', 'prod'}")

# ============================================================
# 6. MODELES DE DONNEES
# ============================================================
doc.add_page_break()
doc.add_heading('6. Modèles de données', level=1)

# --- clients ---
doc.add_heading('App clients (SQLite)', level=2)

doc.add_heading('Utilisateur', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['user', 'OneToOneField(User)', 'Lien vers le modèle Django User'],
        ['code_tiers', 'CharField(20)', 'Code client dans le système ERP'],
    ]
)
add_text("Méthode clé : get_client_distant() — récupère les infos client depuis la BDD MariaDB.")

doc.add_heading('TokenResetPassword', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['user', 'ForeignKey(User)', 'Utilisateur concerné'],
        ['token', 'CharField(64)', 'Token hexadécimal unique'],
        ['date_creation', 'DateTimeField', 'Date de création'],
        ['utilise', 'BooleanField', 'Consommé ou non'],
    ]
)
add_bullet("1 heure", bold_prefix="Validité : ")
add_bullet("secrets.token_hex(32) — cryptographiquement sûr", bold_prefix="Génération : ")
add_bullet("Nettoyage automatique des tokens > 24h via nettoyer_anciens()")

doc.add_heading('UtilisateurSupprime', level=3)
add_text("Corbeille pour les utilisateurs supprimés. Fenêtre de restauration de 5 minutes.")
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['username', 'CharField', 'Nom d\'utilisateur sauvegardé'],
        ['password_hash', 'CharField', 'Hash du mot de passe'],
        ['email', 'EmailField', 'Email sauvegardé'],
        ['code_tiers', 'CharField', 'Code client sauvegardé'],
        ['nom_client', 'CharField', 'Nom du client'],
        ['commandes_json', 'JSONField', 'Commandes au format JSON'],
        ['date_suppression', 'DateTimeField', 'Date de mise en corbeille'],
    ]
)

doc.add_heading('HistoriqueSuppressionUtilisateur', level=3)
add_text("Archive permanente des utilisateurs supprimés définitivement.")

# --- catalogue ---
doc.add_heading('App catalogue (MariaDB — lecture seule)', level=2)
add_text("Tous les modèles ont managed = False (Django ne gère pas les migrations).")

doc.add_heading('Prod (table prod)', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['prod', 'CharField (PK)', 'Code produit'],
        ['libelle', 'CharField', 'Nom du produit'],
    ]
)

doc.add_heading('ComCli (table comcli)', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['tiers', 'IntegerField (PK)', 'Code client'],
        ['nom', 'CharField', 'Nom du client'],
        ['complement', 'CharField', 'Complément d\'adresse'],
        ['adresse', 'CharField', 'Adresse'],
        ['cp', 'CharField', 'Code postal'],
        ['acheminement', 'CharField', 'Ville'],
        ['date_dep', 'DateField', 'Date de départ'],
        ['date_liv', 'DateField', 'Date de livraison'],
    ]
)

doc.add_heading('ComCliLig (table comclilig)', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['prod', 'CharField (PK)', 'Référence produit'],
        ['qte', 'IntegerField', 'Quantité'],
        ['poids', 'DecimalField', 'Poids'],
        ['colis', 'IntegerField', 'Nombre de colis'],
        ['pu_base', 'DecimalField', 'Prix unitaire de base'],
        ['pu_net', 'DecimalField', 'Prix unitaire net'],
        ['commentaire_prep', 'CharField', 'Commentaire préparation'],
    ]
)

doc.add_heading('Catalogue (table catalogue)', level=3)
add_text("Définit quels produits sont disponibles pour quels clients (clé composite tiers + prod).")
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['tiers', 'IntegerField', 'Code client'],
        ['prod', 'CharField', 'Code produit'],
    ]
)

# --- commandes ---
doc.add_page_break()
doc.add_heading('App commandes (SQLite)', level=2)

doc.add_heading('Commande', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['utilisateur', 'ForeignKey(Utilisateur)', 'Client propriétaire'],
        ['numero', 'CharField (unique)', 'Format : CMD-YYYYMMDD-XXXX'],
        ['date_commande', 'DateTimeField (auto)', 'Date de création'],
        ['date_livraison', 'DateField (null)', 'Date de livraison souhaitée'],
        ['date_depart_camions', 'DateField (null)', 'Date de départ des camions'],
        ['total_ht', 'DecimalField', 'Total hors taxes'],
        ['commentaire', 'TextField', 'Commentaire libre'],
    ]
)

doc.add_heading('LigneCommande', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['commande', 'ForeignKey(Commande)', 'Commande parente'],
        ['reference_produit', 'CharField', 'Code produit (snapshot)'],
        ['nom_produit', 'CharField', 'Libellé produit (snapshot)'],
        ['quantite', 'IntegerField', 'Quantité commandée'],
        ['prix_unitaire', 'DecimalField', 'Prix unitaire au moment de la commande'],
        ['total_ligne', 'DecimalField', 'Calcul auto : quantité × prix'],
    ]
)
add_text(
    "Les champs reference_produit, nom_produit et prix_unitaire sont des snapshots "
    "capturés au moment de la commande. Ils ne changent pas si le catalogue évolue.",
    bold=True
)

doc.add_heading('CommandeSupprimee / HistoriqueSuppression', level=3)
add_text(
    "CommandeSupprimee : corbeille avec fenêtre de restauration de 5 minutes. "
    "HistoriqueSuppression : archive permanente des commandes supprimées définitivement."
)

# --- recommandations ---
doc.add_heading('App recommandations (SQLite)', level=2)

doc.add_heading('HistoriqueAchat', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['utilisateur', 'ForeignKey(Utilisateur)', 'Client concerné'],
        ['reference_produit', 'CharField', 'Code produit'],
        ['categorie', 'CharField', 'Catégorie du produit'],
        ['quantite_totale', 'IntegerField', 'Quantité cumulée'],
        ['nombre_commandes', 'IntegerField', 'Nombre de commandes contenant ce produit'],
        ['dernier_achat', 'DateTimeField (auto)', 'Date du dernier achat'],
        ['premier_achat', 'DateTimeField (auto)', 'Date du premier achat'],
    ]
)
add_text("Contrainte unique : (utilisateur, reference_produit)")

doc.add_heading('PreferenceCategorie', level=3)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['utilisateur', 'ForeignKey(Utilisateur)', 'Client concerné'],
        ['categorie', 'CharField', 'Catégorie de produits'],
        ['score', 'FloatField', 'Score de préférence calculé'],
    ]
)

# ============================================================
# 7. ROUTAGE URL
# ============================================================
doc.add_page_break()
doc.add_heading('7. Routage URL', level=1)

doc.add_heading('Routes racine (extranet/urls.py)', level=2)
add_table(
    ['Préfixe', 'Application', 'Namespace'],
    [
        ['/admin/', 'Django admin', '-'],
        ['/', 'clients', 'clients'],
        ['/catalogue/', 'catalogue', 'catalogue'],
        ['/commandes/', 'commandes', 'commandes'],
        ['/recommandations/', 'recommandations', 'recommandations'],
        ['/administration/', 'administration', 'administration'],
    ]
)

doc.add_heading('Routes clients (/)', level=2)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/connexion/', 'connexion', 'Formulaire de connexion'],
        ['/deconnexion/', 'deconnexion', 'Déconnexion'],
        ['/contact/', 'contact', 'Page de contact'],
        ['/profil/', 'profil', 'Profil utilisateur'],
        ['/profil/mot-de-passe/', 'modifier_mot_de_passe', 'Changement de mot de passe'],
        ['/profil/email/', 'modifier_email', 'Changement d\'email'],
        ['/demande-mot-de-passe/', 'demande_mot_de_passe', 'Demande de réinitialisation'],
        ['/reset-password/<token>/', 'reset_password_confirm', 'Réinitialisation avec token'],
    ]
)

doc.add_heading('Routes catalogue (/catalogue/)', level=2)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/catalogue/', 'liste_produits', 'Liste des produits'],
        ['/catalogue/produit/<ref>/', 'detail_produit', 'Détail d\'un produit'],
        ['/catalogue/favoris/', 'favoris', 'Top 4 produits achetés'],
        ['/catalogue/commander/', 'commander', 'Formulaire de commande'],
        ['/catalogue/mentions-legales/', 'mentions_legales', 'Mentions légales'],
    ]
)

doc.add_heading('Routes commandes (/commandes/)', level=2)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/commandes/panier/', 'voir_panier', 'Afficher le panier'],
        ['/commandes/panier/ajouter/', 'ajouter_au_panier', 'Ajouter un produit (POST)'],
        ['/commandes/panier/modifier/', 'modifier_quantite', 'Modifier une quantité (POST)'],
        ['/commandes/panier/supprimer/', 'supprimer_du_panier', 'Retirer un produit (POST)'],
        ['/commandes/panier/vider/', 'vider_panier', 'Vider le panier (POST)'],
        ['/commandes/valider/', 'valider_commande', 'Valider la commande'],
        ['/commandes/confirmation/', 'confirmation_commande', 'Page de confirmation'],
        ['/commandes/historique/', 'historique_commandes', 'Historique (50 dernières)'],
        ['/commandes/details/<id>/', 'details_commande', 'Détails d\'une commande'],
    ]
)

doc.add_heading('Routes administration (/administration/)', level=2)

add_text("Tableau de bord", bold=True)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/administration/', 'dashboard', 'Statistiques et activités'],
    ]
)

add_text("Gestion des commandes", bold=True)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/administration/commandes/', 'liste_commande', 'Liste des commandes'],
        ['/administration/commandes/<id>/', 'details_commande', 'Détails commande'],
        ['/administration/commandes/<id>/supprimer/', 'supprimer_commande', 'Mettre en corbeille'],
        ['…/commandes-supprimees/<id>/restaurer/', 'restaurer_commande', 'Restaurer'],
        ['…/commandes-supprimees/<id>/supprimer/', 'supprimer_definitivement', 'Suppression définitive'],
    ]
)

add_text("Gestion des utilisateurs", bold=True)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/administration/utilisateurs/', 'catalogue_utilisateurs', 'Liste des utilisateurs'],
        ['/administration/information/<id>/', 'information_utilisateur', 'Informations utilisateur'],
        ['…/information/<id>/modifier/', 'modifier_utilisateur', 'Modifier utilisateur'],
        ['…/information/<id>/mot-de-passe/', 'changer_mot_de_passe', 'Changer MDP'],
        ['…/information/<id>/commandes/', 'commande_utilisateur', 'Commandes utilisateur'],
        ['/administration/inscription/', 'inscription', 'Créer un utilisateur'],
    ]
)

add_text("Clients distants et API", bold=True)
add_table(
    ['URL', 'Vue', 'Description'],
    [
        ['/administration/clients/', 'catalogue_clients', 'Liste clients (BDD distante)'],
        ['/administration/cadencier/<code>/', 'cadencier_client', 'Historique prix/commandes'],
        ['/administration/api/recherche-clients/', 'recherche_clients_api', 'Recherche AJAX'],
        ['/administration/api/verifier-mdp/<id>/', 'verifier_mot_de_passe_api', 'Vérification MDP'],
    ]
)

# ============================================================
# 8. VUES ET LOGIQUE METIER
# ============================================================
doc.add_page_break()
doc.add_heading('8. Vues et logique métier', level=1)

doc.add_heading('App clients', level=2)
add_table(
    ['Vue', 'Accès', 'Description'],
    [
        ['connexion', 'Public', 'Authentifie l\'utilisateur. Redirige admin → /administration/, client → /recommandations/'],
        ['deconnexion', 'Authentifié', 'Déconnecte et vide la session'],
        ['profil', 'Authentifié', 'Affiche le profil avec données du système distant'],
        ['modifier_mot_de_passe', 'Authentifié', 'Changement de MDP (ancien MDP requis)'],
        ['modifier_email', 'Authentifié', 'Changement d\'email'],
        ['demande_mot_de_passe', 'Public', 'Crée un TokenResetPassword et envoie un email'],
        ['reset_password_confirm', 'Public', 'Valide le token (1h) et permet un nouveau MDP'],
    ]
)

doc.add_heading('App catalogue', level=2)
add_table(
    ['Vue', 'Accès', 'Description'],
    [
        ['liste_produits', 'Authentifié', 'Liste paginée avec filtres, recherche et favoris'],
        ['detail_produit', 'Authentifié', 'Détail d\'un produit avec prix'],
        ['favoris', 'Authentifié', 'Les 4 produits les plus commandés'],
        ['commander', 'Authentifié', 'Formulaire de commande'],
        ['mentions_legales', 'Public', 'Page des mentions légales'],
    ]
)

doc.add_heading('App commandes', level=2)
add_table(
    ['Vue', 'Accès', 'Description'],
    [
        ['voir_panier', 'Authentifié', 'Contenu du panier (session)'],
        ['ajouter_au_panier', 'Authentifié', 'POST/AJAX : ajoute un produit'],
        ['modifier_quantite', 'Authentifié', 'POST : modifie la quantité'],
        ['supprimer_du_panier', 'Authentifié', 'POST : retire un article'],
        ['vider_panier', 'Authentifié', 'POST : vide le panier'],
        ['valider_commande', 'Authentifié', 'GET : récap / POST : crée commande + EDI + email'],
        ['confirmation_commande', 'Authentifié', 'Numéro de commande et message de succès'],
        ['historique_commandes', 'Authentifié', 'Liste des 50 dernières commandes'],
        ['details_commande', 'Authentifié', 'Détails avec vérification de propriété'],
    ]
)

doc.add_heading('App administration', level=2)
add_text("Toutes les vues sont protégées par le décorateur @admin_required (is_staff=True requis).", bold=True)
add_table(
    ['Vue', 'Description'],
    [
        ['dashboard', 'Statistiques, activités récentes, corbeilles'],
        ['liste_commande', 'Liste filtrable de toutes les commandes'],
        ['details_commande', 'Détail avec lignes de commande'],
        ['supprimer/restaurer_commande', 'Suppression douce et restauration'],
        ['catalogue_utilisateurs', 'Liste de tous les utilisateurs extranet'],
        ['information_utilisateur', 'Détail utilisateur avec infos distantes'],
        ['modifier_utilisateur', 'Édition code_tiers et informations'],
        ['inscription', 'Création d\'un nouvel utilisateur'],
        ['catalogue_clients', 'Liste clients depuis la BDD distante'],
        ['cadencier_client', 'Historique prix/commandes d\'un client'],
    ]
)

# ============================================================
# 9. TEMPLATES
# ============================================================
doc.add_page_break()
doc.add_heading('9. Templates', level=1)

doc.add_heading('Organisation', level=2)
add_text("Les templates sont séparés en deux espaces :")

doc.add_heading('Côté client (templates/cote_client/)', level=3)
add_table(
    ['Fichier', 'Description'],
    [
        ['base.html', 'Layout principal (navbar, footer, messages)'],
        ['navbar.html', 'Barre de navigation avec menu utilisateur et panier'],
        ['catalogue/liste.html', 'Grille de produits avec filtres'],
        ['catalogue/detail.html', 'Fiche produit'],
        ['catalogue/favoris.html', 'Produits favoris'],
        ['catalogue/commander.html', 'Formulaire de commande'],
        ['catalogue/recap_panier.html', 'Récapitulatif du panier'],
        ['catalogue/filtres.html', 'Barre latérale de filtres'],
    ]
)

doc.add_heading('Côté administration (templates/administration/)', level=3)
add_table(
    ['Fichier', 'Description'],
    [
        ['base_admin.html', 'Layout admin (sidebar, navbar, footer)'],
        ['dashboard.html', 'Tableau de bord'],
        ['includes/navbar.html', 'Barre supérieure admin'],
        ['includes/sidebar.html', 'Menu latéral'],
        ['includes/_cartes_statistiques.html', 'Cartes de statistiques'],
        ['commandes/liste_commandes.html', 'Liste des commandes'],
        ['commandes/details_commande.html', 'Détails commande'],
        ['utilisateurs/catalogue_utilisateur.html', 'Liste des utilisateurs'],
        ['utilisateurs/information_utilisateur.html', 'Détail utilisateur'],
        ['clients/catalogue_clients.html', 'Liste clients distants'],
        ['clients/cadencier_client.html', 'Historique commandes client'],
    ]
)

doc.add_heading('Héritage des templates', level=3)
add_code_block(
    "base.html (côté client)\n"
    "  └── navbar.html\n"
    "  └── {% block content %}\n"
    "  └── Messages Django\n"
    "\n"
    "base_admin.html (administration)\n"
    "  └── includes/sidebar.html\n"
    "  └── includes/navbar.html\n"
    "  └── {% block content %}\n"
    "  └── Messages Django"
)

doc.add_heading('Context processors globaux', level=3)
add_table(
    ['Processeur', 'Variable', 'Description'],
    [
        ['commandes.context_processors.panier_count', 'panier_count', 'Nombre d\'articles dans le panier'],
    ]
)

# ============================================================
# 10. FICHIERS STATIQUES
# ============================================================
doc.add_page_break()
doc.add_heading('10. Fichiers statiques', level=1)

doc.add_heading('CSS', level=2)
add_table(
    ['Fichier', 'Description'],
    [
        ['css/style.css', 'Styles principaux de l\'application'],
        ['css/style_admin.css', 'Styles spécifiques à l\'administration'],
    ]
)

doc.add_heading('JavaScript', level=2)
add_table(
    ['Fichier', 'Description'],
    [
        ['js/panier.js', 'Gestion du panier (ajout, suppression, AJAX)'],
        ['js/accordion-filtres.js', 'Accordéon pour le panneau de filtres'],
        ['js/client-search.js', 'Recherche de clients en temps réel (AJAX)'],
        ['js/commander.js', 'Logique du formulaire de commande'],
        ['js/countdown.js', 'Affichage de compte à rebours (corbeille)'],
        ['js/date-calcul.js', 'Calculs de dates (livraison, départ camions)'],
        ['js/liste_limite.js', 'Gestion de la pagination/limite des listes'],
        ['js/retour-haut.js', 'Bouton de retour en haut de page'],
        ['js/toggle-password.js', 'Bascule visibilité du mot de passe'],
    ]
)

doc.add_heading('Bibliothèques externes (CDN)', level=2)
add_table(
    ['Bibliothèque', 'Version', 'Usage'],
    [
        ['Bootstrap CSS', '5.3.2', 'Framework CSS responsive'],
        ['Bootstrap JS Bundle', '5.3.2', 'Composants interactifs'],
        ['Bootstrap Icons', '1.11.1', 'Icônes'],
    ]
)

# ============================================================
# 11. AUTHENTIFICATION ET AUTORISATION
# ============================================================
doc.add_page_break()
doc.add_heading('11. Authentification et autorisation', level=1)

doc.add_heading('Niveaux d\'accès', level=2)
add_table(
    ['Niveau', 'Décorateur', 'Routes accessibles'],
    [
        ['Public', '-', '/connexion, /demande-mot-de-passe, /reset-password, /mentions-legales'],
        ['Client connecté', '@login_required', '/catalogue/*, /commandes/*, /profil/*, /recommandations/*'],
        ['Administrateur', '@admin_required', '/administration/* (is_staff=True requis)'],
    ]
)

doc.add_heading('Flux de connexion', level=2)
add_bullet("L'utilisateur soumet le formulaire /connexion/")
add_bullet("Django authentifie via authenticate() et login()")
add_bullet("Vérification de l'existence du profil Utilisateur")
add_bullet("Redirection selon le rôle :")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(2)
run = p.add_run("is_staff=True → /administration/\nSinon → /recommandations/")
run.font.size = Pt(10)

doc.add_heading('Décorateur @admin_required', level=2)
add_text(
    "Défini dans administration/views/utils/decorators.py. "
    "Vérifie is_authenticated et is_staff. Redirige vers la page de connexion si non autorisé."
)

doc.add_heading('Flux de réinitialisation du mot de passe', level=2)
add_bullet("L'utilisateur visite /demande-mot-de-passe/")
add_bullet("Création d'un TokenResetPassword (validité 1 heure)")
add_bullet("Envoi d'un email avec lien /reset-password/<token>/")
add_bullet("L'utilisateur clique et soumet un nouveau mot de passe")
add_bullet("Le token est marqué comme utilisé")
add_bullet("Nettoyage automatique des tokens > 24h")

doc.add_heading('Validateurs de mot de passe', level=2)
add_bullet("Empêche les MDP similaires aux attributs utilisateur", bold_prefix="UserAttributeSimilarityValidator : ")
add_bullet("Minimum 8 caractères", bold_prefix="MinimumLengthValidator : ")
add_bullet("Bloque les MDP courants", bold_prefix="CommonPasswordValidator : ")
add_bullet("Empêche les MDP entièrement numériques", bold_prefix="NumericPasswordValidator : ")

# ============================================================
# 12. POINTS D'ENTREE API
# ============================================================
doc.add_page_break()
doc.add_heading('12. Points d\'entrée API', level=1)

doc.add_heading('API JSON Recommandations', level=2)

doc.add_heading('GET /recommandations/api/', level=3)
add_text("Retourne les produits recommandés pour l'utilisateur connecté.")
add_table(
    ['Paramètre', 'Type', 'Défaut', 'Description'],
    [
        ['limite', 'int', '10', 'Nombre max de résultats'],
    ]
)
add_text("Réponse :")
add_code_block(
    '{\n'
    '  "recommandations": [\n'
    '    {"reference": "PROD001", "nom": "Nom du produit", "prix": 12.50}\n'
    '  ]\n'
    '}'
)

doc.add_heading('GET /recommandations/api/favoris/', level=3)
add_text("Retourne les produits favoris (les plus commandés).")
add_table(
    ['Paramètre', 'Type', 'Défaut', 'Description'],
    [
        ['limite', 'int', '4', 'Nombre max de résultats'],
    ]
)

doc.add_heading('API Administration (AJAX)', level=2)

doc.add_heading('GET /administration/api/recherche-clients/', level=3)
add_text("Recherche de clients par nom dans la base distante.")
add_table(
    ['Paramètre', 'Type', 'Description'],
    [
        ['q', 'string', 'Terme de recherche'],
    ]
)

doc.add_heading('GET /administration/api/verifier-mdp/<id>/', level=3)
add_text("Vérifie un mot de passe avant une opération sensible.")
add_table(
    ['Paramètre', 'Type', 'Description'],
    [
        ['password', 'string', 'Mot de passe à vérifier'],
    ]
)
add_text("Réponse :")
add_code_block('{"valid": true}')

doc.add_heading('Opérations Panier (POST/AJAX)', level=2)
add_table(
    ['Endpoint', 'Paramètres POST', 'Description'],
    [
        ['/commandes/panier/ajouter/', 'reference, quantite', 'Ajouter au panier'],
        ['/commandes/panier/modifier/', 'reference, quantite', 'Modifier la quantité'],
        ['/commandes/panier/supprimer/', 'reference', 'Retirer du panier'],
        ['/commandes/panier/vider/', '-', 'Vider tout le panier'],
    ]
)

# ============================================================
# 13. SYSTEME DE RECOMMANDATIONS
# ============================================================
doc.add_page_break()
doc.add_heading('13. Système de recommandations', level=1)

doc.add_heading('Fonctionnement', level=2)
add_text("Le système de recommandations est basé sur l'historique d'achats de chaque utilisateur.")

doc.add_heading('Cycle de vie des données', level=2)
add_bullet("Validation d'une commande")
add_bullet("HistoriqueAchat.enregistrer_achat() — incrémente quantité et nombre de commandes")
add_bullet("Calcul des PreferenceCategorie — score basé sur fréquence et récence")
add_bullet("Affichage sur /recommandations/ — produits triés par score")

doc.add_heading('Points d\'accès', level=2)
add_table(
    ['Type', 'URL', 'Description'],
    [
        ['Page', '/recommandations/', 'Interface visuelle'],
        ['API', '/recommandations/api/', 'Données JSON'],
        ['API Favoris', '/recommandations/api/favoris/', 'Top 4 produits'],
    ]
)

# ============================================================
# 14. GESTION DU PANIER
# ============================================================
doc.add_heading('14. Gestion du panier', level=1)

doc.add_heading('Stockage en session', level=2)
add_text("Le panier est stocké dans request.session['panier'] sous forme de dictionnaire :")
add_code_block(
    "session['panier'] = {\n"
    "    'PROD001': {\n"
    "        'reference': 'PROD001',\n"
    "        'nom': 'Nom du produit',\n"
    "        'quantite': 3,\n"
    "        'prix_unitaire': 12.50\n"
    "    }\n"
    "}"
)

doc.add_heading('Avantages', level=2)
add_bullet("Pas de table supplémentaire en base de données")
add_bullet("Panier automatiquement vidé à la fermeture du navigateur")
add_bullet("Performances optimales (pas de requêtes BDD pour le panier)")

doc.add_heading('Flux de commande', level=2)
add_bullet("Ajout de produits au panier (session)")
add_bullet("Consultation du panier (/commandes/panier/)")
add_bullet("Validation (/commandes/valider/) :")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(2)
run = p.add_run(
    "— Crée un objet Commande (numéro auto-généré)\n"
    "— Crée les LigneCommande (snapshot des prix)\n"
    "— Génère le fichier EDI CSV\n"
    "— Envoie un email de confirmation\n"
    "— Met à jour l'HistoriqueAchat\n"
    "— Vide le panier en session"
)
run.font.size = Pt(10)
add_bullet("Confirmation (/commandes/confirmation/)")

# ============================================================
# 15. EXPORT EDI
# ============================================================
doc.add_page_break()
doc.add_heading('15. Export EDI', level=1)

doc.add_heading('Description', level=2)
add_text(
    "À chaque validation de commande, un fichier CSV au format EDI (ORDERS) est généré "
    "pour intégration avec le système ERP."
)

doc.add_heading('Format du fichier', level=2)
add_text("Le fichier suit une structure inspirée du standard UN/EDIFACT :")
add_code_block(
    "@GP;...                    # En-tête générale\n"
    "ENT;CMD-20260212-0001;...  # Entité commande\n"
    "DTM;20260212;20260214;...  # Dates (commande, livraison, départ)\n"
    "PAR;...                    # Parties (fournisseur, acheteur)\n"
    "LIG;PROD001;3;12.50;...   # Ligne de commande\n"
    "END;...                    # Fin de document"
)

doc.add_heading('Emplacement des fichiers', level=2)
add_code_block("edi_exports/\n  ├── CMD-20260212-0001.csv\n  ├── CMD-20260212-0002.csv\n  └── ...")

doc.add_heading('Fonction principale', level=2)
add_text(
    "generer_csv_edi(commande, client, lignes) dans commandes/services.py : "
    "reçoit l'objet commande, les infos client et les lignes, "
    "génère le fichier CSV et le sauvegarde dans EDI_OUTPUT_DIR."
)

# ============================================================
# 16. SUPPRESSION DOUCE ET RESTAURATION
# ============================================================
doc.add_heading('16. Suppression douce et restauration', level=1)

doc.add_heading('Principe', level=2)
add_text(
    "Le système implémente un mécanisme de corbeille avec une fenêtre de restauration "
    "de 5 minutes pour les commandes et les utilisateurs."
)

doc.add_heading('Flux de suppression', level=2)
add_code_block(
    "Élément actif\n"
    "     |\n"
    "     v\n"
    "[Suppression douce]  -->  Corbeille (5 min)\n"
    "                              |\n"
    "                  +-----------+-----------+\n"
    "                  |                       |\n"
    "           [Restauration]        [Expiration 5 min]\n"
    "                  |                       |\n"
    "                  v                       v\n"
    "           Élément restauré       [Suppression définitive]\n"
    "                                          |\n"
    "                                          v\n"
    "                                   Archive permanente"
)

doc.add_heading('Commandes', level=3)
add_table(
    ['Étape', 'Table cible', 'Description'],
    [
        ['Suppression douce', 'CommandeSupprimee', 'Données + lignes sauvegardées en JSON'],
        ['Restauration', 'Commande + LigneCommande', 'Recrée la commande et ses lignes'],
        ['Suppression définitive', 'HistoriqueSuppression', 'Archive numéro, client, total, date'],
    ]
)

doc.add_heading('Utilisateurs', level=3)
add_table(
    ['Étape', 'Table cible', 'Description'],
    [
        ['Suppression douce', 'UtilisateurSupprime', 'Données + commandes en JSON'],
        ['Restauration', 'User + Utilisateur + Commande', 'Recrée l\'utilisateur et ses commandes'],
        ['Suppression définitive', 'HistoriqueSuppressionUtilisateur', 'Archive permanente'],
    ]
)

doc.add_heading('Méthodes utilitaires', level=3)
add_bullet("Vérifie si la fenêtre de 5 minutes n'est pas expirée", bold_prefix="est_restaurable() : ")
add_bullet("Retourne le temps restant avant expiration", bold_prefix="temps_restant() : ")
add_bullet("Supprime les éléments expirés de la corbeille", bold_prefix="nettoyer_anciens() : ")

# ============================================================
# 17. SYSTEME DE FILTRES
# ============================================================
doc.add_page_break()
doc.add_heading('17. Système de filtres', level=1)

doc.add_heading('Architecture', level=2)
add_text("Le système de filtres est défini dans deux emplacements :")
add_bullet("Filtres pour le catalogue client", bold_prefix="catalogue/services.py : ")
add_bullet("Filtres pour l'administration", bold_prefix="administration/views/utils/filtres.py : ")

doc.add_heading('Catégories de filtres', level=2)
add_text("Le système propose plus de 30 catégories organisées en groupes :")
add_table(
    ['Groupe', 'Exemples de filtres'],
    [
        ['Format', 'Frais, Surgelé, Sous-vide'],
        ['Viandes', 'Bœuf, Veau, Porc, Agneau, Volaille'],
        ['Découpes', 'Filet, Côte, Entrecôte, Rôti, Steak'],
        ['Charcuterie', 'Saucisse, Jambon, Pâté, Terrine, Boudin'],
        ['Préparations', 'Haché, Brochette, Marinade, Pané'],
        ['Fromages', 'Fromage, Comté, Emmental, Camembert'],
        ['Autres', 'Œuf, Beurre, Crème, Sauce'],
    ]
)

doc.add_heading('Fonctionnement', level=2)
add_bullet("Analyse la liste de produits et génère les filtres applicables", bold_prefix="preparer_filtres(produits) : ")
add_bullet("Filtre la liste selon les filtres sélectionnés", bold_prefix="appliquer_filtres(produits, filtres_actifs) : ")
add_bullet("Crée des filtres à partir des libellés (min. 3 occurrences)", bold_prefix="generer_filtres_automatiques() : ")
add_bullet("Normalisation Unicode (NFD) pour comparaisons insensibles aux accents", bold_prefix="_normaliser(texte) : ")

# ============================================================
# 18. DEPENDANCES
# ============================================================
doc.add_heading('18. Dépendances', level=1)

doc.add_heading('Backend (requirements.txt)', level=2)
add_table(
    ['Package', 'Version', 'Description'],
    [
        ['Django', '6.0.1', 'Framework web'],
        ['asgiref', '3.11.0', 'Utilitaires ASGI'],
        ['sqlparse', '0.5.5', 'Parsing SQL'],
        ['tzdata', '2025.3', 'Base de données des fuseaux horaires'],
    ]
)

doc.add_heading('Frontend (CDN)', level=2)
add_table(
    ['Bibliothèque', 'Version', 'Usage'],
    [
        ['Bootstrap CSS/JS', '5.3.2', 'Framework CSS responsive + composants'],
        ['Bootstrap Icons', '1.11.1', 'Icônes'],
    ]
)

# ============================================================
# 19. VARIABLES D'ENVIRONNEMENT
# ============================================================
doc.add_page_break()
doc.add_heading('19. Variables d\'environnement', level=1)

add_text("Le projet utilise un fichier .env (non versionné) pour les configurations sensibles.")

doc.add_heading('Variables requises', level=2)
add_table(
    ['Variable', 'Description', 'Exemple'],
    [
        ['SECRET_KEY', 'Clé secrète Django', 'django-insecure-xxx...'],
        ['DEBUG', 'Mode debug', 'True / False'],
        ['ALLOWED_HOSTS', 'Hôtes autorisés', 'localhost,127.0.0.1'],
    ]
)

doc.add_heading('Base de données distante', level=2)
add_table(
    ['Variable', 'Description', 'Exemple'],
    [
        ['DB_LOGIGVD_NAME', 'Nom de la base MariaDB', 'logigvd'],
        ['DB_LOGIGVD_USER', 'Utilisateur MariaDB', 'extranet_user'],
        ['DB_LOGIGVD_PASSWORD', 'Mot de passe MariaDB', '********'],
        ['DB_LOGIGVD_HOST', 'Hôte du serveur MariaDB', '192.168.1.100'],
        ['DB_LOGIGVD_PORT', 'Port MariaDB', '3306'],
    ]
)


# ============================================================
# 20. DEPLOIEMENT
# ============================================================
doc.add_heading('20. Déploiement', level=1)

doc.add_heading('Prérequis', level=2)
add_bullet("Python 3.10+")
add_bullet("Accès réseau au serveur MariaDB (BDD distante)")
add_bullet("Serveur SMTP ou backend console pour les emails (optionnel en dev)")

doc.add_heading('Installation', level=2)
add_code_block(
    "# 1. Cloner le dépôt\n"
    "git clone <url-du-depot>\n"
    "cd Code\n"
    "\n"
    "# 2. Créer l'environnement virtuel\n"
    "python -m venv env\n"
    "env\\Scripts\\activate          # Windows\n"
    "source env/bin/activate        # Linux/Mac\n"
    "\n"
    "# 3. Installer les dépendances\n"
    "pip install -r requirements.txt\n"
    "\n"
    "# 4. Configurer les variables d'environnement\n"
    "cp .env.example .env\n"
    "# Éditer .env avec les valeurs appropriées\n"
    "\n"
    "# 5. Appliquer les migrations (BDD locale uniquement)\n"
    "python manage.py migrate\n"
    "\n"
    "# 6. Créer un superutilisateur\n"
    "python manage.py createsuperuser\n"
    "\n"
    "# 7. Lancer le serveur de développement\n"
    "python manage.py runserver"
)

doc.add_heading('Production', level=2)

add_text("Points d'entrée serveur", bold=True)
add_table(
    ['Fichier', 'Protocole', 'Usage'],
    [
        ['extranet/wsgi.py', 'WSGI', 'Gunicorn, Apache mod_wsgi'],
        ['extranet/asgi.py', 'ASGI', 'Daphne, Uvicorn (support async)'],
    ]
)

add_text("Checklist de déploiement", bold=True)
add_bullet("DEBUG = False dans .env")
add_bullet("SECRET_KEY unique et sécurisée")
add_bullet("ALLOWED_HOSTS configuré avec les domaines de production")
add_bullet("Configurer STATIC_ROOT et exécuter python manage.py collectstatic")
add_bullet("Configurer l'envoi d'emails si nécessaire")
add_bullet("Configurer le dossier edi_exports/ avec les permissions adéquates")
add_bullet("Mettre en place un reverse proxy (Nginx) devant le serveur applicatif")
add_bullet("Configurer HTTPS avec un certificat SSL")
add_bullet("Mettre en place les sauvegardes automatiques de la BDD SQLite")
add_bullet("Configurer CSRF_TRUSTED_ORIGINS avec les domaines de production")

add_text("Exemple de configuration Gunicorn", bold=True)
add_code_block(
    "gunicorn extranet.wsgi:application \\\n"
    "    --bind 0.0.0.0:8000 \\\n"
    "    --workers 3 \\\n"
    "    --timeout 120"
)

# ============================================================
# ANNEXE
# ============================================================
doc.add_page_break()
doc.add_heading('Annexe : Diagramme des relations entre modèles', level=1)

add_code_block(
    "                    +-------------------+\n"
    "                    |    User (Django)  |\n"
    "                    +-------------------+\n"
    "                           | 1:1\n"
    "                           v\n"
    "                    +-------------------+\n"
    "                    |   Utilisateur     |\n"
    "                    |   (code_tiers)    |\n"
    "                    +-------------------+\n"
    "                     /        |        \\\n"
    "                    /         |         \\\n"
    "                   v          v          v\n"
    "        +-----------+  +-----------+  +-------------------+\n"
    "        | Commande  |  |Historique |  |TokenResetPassword |\n"
    "        | (numero)  |  |  Achat    |  +-------------------+\n"
    "        +-----------+  +-----------+\n"
    "             |                         +-------------------+\n"
    "             v                         | Preference        |\n"
    "        +-----------+                  | Categorie         |\n"
    "        |  Ligne    |                  +-------------------+\n"
    "        | Commande  |\n"
    "        +-----------+\n"
    "\n"
    "   --- Corbeille ---\n"
    "   UtilisateurSupprime  -->  HistoriqueSuppressionUtilisateur\n"
    "   CommandeSupprimee    -->  HistoriqueSuppression\n"
    "\n"
    "   --- BDD Distante (MariaDB, lecture seule) ---\n"
    "   +--------+    +-----------+    +----------+\n"
    "   |  Prod  |<-->| Catalogue |<-->|  ComCli  |\n"
    "   +--------+    +-----------+    +----------+\n"
    "                                       |\n"
    "                                       v\n"
    "                                  +-----------+\n"
    "                                  | ComCliLig |\n"
    "                                  +-----------+"
)

# ============================================================
# PIED DE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Documentation générée le 12/02/2026 — Extranet Giffaud Groupe')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ============================================================
# SAUVEGARDE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'DOCUMENTATION_TECHNIQUE.docx')
doc.save(output_path)
print(f"Document Word généré avec succès : {output_path}")
