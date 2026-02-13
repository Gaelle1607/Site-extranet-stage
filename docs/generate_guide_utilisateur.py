"""
Script de generation du guide utilisateur au format Word (.docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): '1a3c6e', qn('w:val'): 'clear'})
        shading.append(bg)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        if r % 2 == 0:
            for c in range(len(headers)):
                shading = table.rows[r + 1].cells[c]._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'f0f4f8', qn('w:val'): 'clear'})
                shading.append(bg)
    doc.add_paragraph()
    return table


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_text(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    p.add_run(text)
    return p


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('Note : ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xc0, 0x7b, 0x00)
    p.add_run(text)
    return p


# ============================================================
# PAGE DE GARDE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Guide Utilisateur')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Extranet Giffaud Groupe')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Plateforme de commande en ligne')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

for _ in range(6):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('12 février 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE DES MATIERES
# ============================================================
doc.add_heading('Table des matières', level=1)

toc_items = [
    '1. Se connecter',
    '2. Naviguer dans l\'application',
    '3. Consulter le catalogue',
    '4. Utiliser les filtres',
    '5. Gérer son panier',
    '6. Passer commande via le catalogue',
    '7. Passer commande via le formulaire rapide',
    '8. Valider et confirmer une commande',
    '9. Consulter l\'historique des commandes',
    '10. Gérer son profil',
    '11. Réinitialiser son mot de passe',
    '12. Aide et contact',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. SE CONNECTER
# ============================================================
doc.add_heading('1. Se connecter', level=1)

add_text(
    "Pour accéder à l'Extranet Giffaud Groupe, vous devez disposer d'identifiants "
    "fournis par votre commercial."
)

doc.add_heading('Procédure de connexion', level=2)

add_step(1, 'Ouvrez votre navigateur et accédez à l\'adresse de l\'Extranet.')
add_step(2, 'Saisissez votre identifiant dans le champ "Identifiant".')
add_step(3, 'Saisissez votre mot de passe dans le champ "Mot de passe".')
add_step(4, 'Cliquez sur le bouton "Se connecter".')

add_note(
    'Vous pouvez cliquer sur l\'icône en forme d\'œil à droite du champ mot de passe '
    'pour afficher ou masquer votre mot de passe.'
)

doc.add_heading('Après la connexion', level=2)
add_text(
    'Une fois connecté, vous êtes redirigé vers la page de recommandations personnalisées. '
    'Un message de bienvenue s\'affiche avec le nom de votre entreprise.'
)

doc.add_heading('En cas de problème', level=2)
add_bullet(
    'Vérifiez que votre identifiant et votre mot de passe sont corrects.',
)
add_bullet(
    'Si vous avez oublié votre mot de passe, consultez la section 11 de ce guide.',
)
add_bullet(
    'Si vous ne disposez pas d\'identifiants, contactez votre commercial Giffaud Groupe.',
)

# ============================================================
# 2. NAVIGUER DANS L'APPLICATION
# ============================================================
doc.add_page_break()
doc.add_heading('2. Naviguer dans l\'application', level=1)

add_text(
    'L\'application dispose d\'une barre de navigation en haut de chaque page qui vous '
    'permet d\'accéder rapidement à toutes les fonctionnalités.'
)

doc.add_heading('Menu principal', level=2)
add_table(
    ['Élément', 'Description'],
    [
        ['GIFFAUD GROUPE (logo)', 'Retour à la page d\'accueil du catalogue'],
        ['Recommandations', 'Vos recommandations personnalisées basées sur vos achats'],
        ['Catalogue', 'Consulter l\'ensemble des produits disponibles'],
        ['Favoris', 'Vos produits les plus commandés (top 4)'],
        ['Commander', 'Formulaire de commande rapide (tableau)'],
    ]
)

doc.add_heading('Menu de droite', level=2)
add_table(
    ['Élément', 'Description'],
    [
        ['Panier (icône)', 'Accéder au panier — un badge rouge indique le nombre d\'articles'],
        ['Mes commandes', 'Consulter l\'historique de vos commandes'],
        ['Votre nom', 'Accéder à votre profil'],
        ['Déconnexion', 'Se déconnecter de l\'application'],
    ]
)

doc.add_heading('Éléments présents sur chaque page', level=2)
add_bullet('Messages d\'information, de succès ou d\'erreur en haut de la page (fermables).')
add_bullet('Bouton "Retour en haut" en bas à droite pour remonter rapidement.')
add_bullet('Pied de page avec liens vers "Contactez-nous" et "Mentions légales".')

add_note(
    'Sur mobile et tablette, le menu se replie dans un bouton "hamburger" (trois barres horizontales). '
    'Appuyez dessus pour ouvrir le menu.'
)

# ============================================================
# 3. CONSULTER LE CATALOGUE
# ============================================================
doc.add_page_break()
doc.add_heading('3. Consulter le catalogue', level=1)

add_text(
    'Le catalogue affiche l\'ensemble des produits disponibles pour votre compte client. '
    'Les prix affichés sont en HT (hors taxes) et sont spécifiques à votre entreprise.'
)

doc.add_heading('Présentation de la page', level=2)
add_text('La page catalogue est organisée en trois colonnes :')
add_bullet('la barre de filtres sur le côté gauche,', bold_prefix='Colonne gauche : ')
add_bullet('la grille de produits au centre,', bold_prefix='Colonne centrale : ')
add_bullet('le résumé du panier sur le côté droit.', bold_prefix='Colonne droite : ')

doc.add_heading('Informations affichées par produit', level=2)
add_text('Chaque produit est présenté sous forme de carte avec :')
add_bullet('la référence du produit (en gris),')
add_bullet('le nom du produit,')
add_bullet('le prix unitaire HT et l\'unité (kg, pièce, etc.),')
add_bullet('un champ quantité avec boutons + et - pour ajuster,')
add_bullet('un bouton "Ajouter" pour mettre le produit dans le panier.')

doc.add_heading('Ajouter un produit au panier depuis le catalogue', level=2)
add_step(1, 'Repérez le produit souhaité dans la grille.')
add_step(2, 'Ajustez la quantité avec les boutons + et - (ou saisissez-la directement).')
add_step(3, 'Cliquez sur "Ajouter".')
add_step(4, 'Un message de confirmation apparaît et le compteur du panier se met à jour.')

doc.add_heading('Consulter le détail d\'un produit', level=2)
add_text(
    'En cliquant sur le nom d\'un produit, vous accédez à sa fiche détaillée. '
    'Celle-ci affiche la référence, le nom complet, les catégories, le prix '
    'et une description lorsqu\'elle est disponible. Vous pouvez également '
    'ajouter le produit au panier depuis cette page.'
)

doc.add_heading('Page des favoris', level=2)
add_text(
    'La page "Favoris" accessible depuis le menu affiche vos 4 produits les plus '
    'commandés. Elle vous permet de les retrouver et de les ajouter rapidement au panier.'
)

# ============================================================
# 4. UTILISER LES FILTRES
# ============================================================
doc.add_page_break()
doc.add_heading('4. Utiliser les filtres', level=1)

add_text(
    'Le panneau de filtres, situé à gauche du catalogue, vous permet de retrouver '
    'rapidement les produits qui vous intéressent.'
)

doc.add_heading('Recherche par texte', level=2)
add_text(
    'En haut du panneau de filtres se trouve un champ de recherche. '
    'Saisissez un nom de produit ou une référence pour filtrer la liste.'
)

doc.add_heading('Filtres par catégorie', level=2)
add_text(
    'Les filtres sont regroupés en catégories que vous pouvez ouvrir/fermer en cliquant dessus '
    '(système d\'accordéon). Exemples de catégories :'
)
add_table(
    ['Catégorie', 'Exemples'],
    [
        ['Format', 'Frais, Surgelé, Sous-vide'],
        ['Viandes', 'Bœuf, Veau, Porc, Agneau, Volaille'],
        ['Découpes', 'Filet, Côte, Entrecôte, Rôti, Steak'],
        ['Charcuterie', 'Saucisse, Jambon, Pâté, Terrine'],
        ['Préparations', 'Haché, Brochette, Marinade, Pané'],
        ['Fromages', 'Comté, Emmental, Camembert'],
    ]
)

doc.add_heading('Appliquer et retirer les filtres', level=2)
add_step(1, 'Ouvrez une catégorie en cliquant sur son nom.')
add_step(2, 'Cochez un ou plusieurs filtres.')
add_step(3, 'Cliquez sur le bouton "Filtrer" pour appliquer.')
add_step(4, 'Pour retirer tous les filtres, cliquez sur "Enlever les filtres".')

add_note(
    'Le nombre de produits correspondant à vos filtres est affiché en haut du catalogue. '
    'Si aucun produit ne correspond, un message vous invite à modifier vos critères.'
)

# ============================================================
# 5. GERER SON PANIER
# ============================================================
doc.add_page_break()
doc.add_heading('5. Gérer son panier', level=1)

doc.add_heading('Aperçu rapide (colonne droite du catalogue)', level=2)
add_text(
    'Lorsque vous naviguez dans le catalogue, un résumé de votre panier est visible '
    'sur le côté droit. Il affiche :'
)
add_bullet('la liste des produits ajoutés avec leurs quantités et prix,')
add_bullet('le total HT,')
add_bullet('un bouton "Modifier" pour accéder au panier détaillé,')
add_bullet('un bouton "Commander" pour passer directement à la validation.')

doc.add_heading('Page du panier détaillé', level=2)
add_text(
    'Cliquez sur l\'icône panier dans la barre de navigation ou sur le bouton "Modifier" '
    'pour accéder à la page complète du panier.'
)

add_text('Sur cette page, vous pouvez :', bold=True)
add_bullet('modifier la quantité de chaque article (les totaux se recalculent automatiquement),')
add_bullet('supprimer un article en cliquant sur l\'icône corbeille,')
add_bullet('vider entièrement le panier avec le bouton "Vider le panier",')
add_bullet('continuer vos achats avec le bouton "Continuer mes achats".')

doc.add_heading('Informations de livraison', level=2)
add_text(
    'Sur le côté droit de la page panier, vous trouverez :'
)
add_bullet(
    'calculée automatiquement (champ en lecture seule).',
    bold_prefix='Date de départ camions : '
)
add_bullet(
    'à sélectionner obligatoirement avec le calendrier.',
    bold_prefix='Date de livraison : '
)
add_bullet(
    'un champ libre pour ajouter des instructions à votre commande.',
    bold_prefix='Commentaires : '
)

add_note(
    'Le panier est sauvegardé dans votre session. Il est automatiquement vidé '
    'si vous fermez votre navigateur.'
)

# ============================================================
# 6. PASSER COMMANDE VIA LE CATALOGUE
# ============================================================
doc.add_page_break()
doc.add_heading('6. Passer commande via le catalogue', level=1)

add_text(
    'La méthode la plus courante pour passer commande consiste à ajouter des produits '
    'depuis le catalogue, puis à valider votre panier.'
)

doc.add_heading('Étapes', level=2)
add_step(1, 'Parcourez le catalogue et ajoutez les produits souhaités au panier.')
add_step(2, 'Cliquez sur l\'icône panier ou "Commander" dans le résumé latéral.')
add_step(3, 'Vérifiez les quantités et ajustez si nécessaire.')
add_step(4, 'Sélectionnez une date de livraison.')
add_step(5, 'Ajoutez un commentaire si besoin.')
add_step(6, 'Cliquez sur "Suivant" pour accéder au récapitulatif.')
add_step(7, 'Vérifiez le récapitulatif et cliquez sur "Confirmer la commande".')

# ============================================================
# 7. PASSER COMMANDE VIA LE FORMULAIRE RAPIDE
# ============================================================
doc.add_page_break()
doc.add_heading('7. Passer commande via le formulaire rapide', level=1)

add_text(
    'Le formulaire rapide ("Commander" dans le menu) est conçu pour passer des commandes '
    'en volume. Tous les produits sont affichés dans un tableau unique.'
)

doc.add_heading('Présentation', level=2)
add_text('La page est organisée en deux parties :')
add_bullet(
    'un tableau avec tous vos produits, leur prix, un champ quantité et le sous-total.',
    bold_prefix='Section gauche : '
)
add_bullet(
    'vos informations client, le récapitulatif, la date de livraison et les commentaires.',
    bold_prefix='Section droite : '
)

doc.add_heading('Utilisation', level=2)
add_step(1, 'Utilisez la barre de recherche pour trouver un produit précis.')
add_step(2, 'Utilisez le bouton "Filtres" pour filtrer par catégorie.')
add_step(3, 'Saisissez la quantité souhaitée pour chaque produit à commander.')
add_step(4, 'Les sous-totaux et le total se mettent à jour en temps réel.')
add_step(5, 'Sélectionnez une date de livraison (obligatoire).')
add_step(6, 'Ajoutez un commentaire si nécessaire.')
add_step(7, 'Cliquez sur "Récapitulatif de la commande" pour valider.')

add_note(
    'Les produits dont la quantité est à 0 ne seront pas inclus dans la commande.'
)

# ============================================================
# 8. VALIDER ET CONFIRMER UNE COMMANDE
# ============================================================
doc.add_page_break()
doc.add_heading('8. Valider et confirmer une commande', level=1)

doc.add_heading('Page de récapitulatif', level=2)
add_text(
    'Avant la confirmation finale, une page de récapitulatif vous permet de vérifier '
    'l\'ensemble de votre commande :'
)
add_bullet('la liste de tous les articles avec quantités et prix,')
add_bullet('le total HT de la commande,')
add_bullet('votre nom de client,')
add_bullet('la date de livraison et la date de départ des camions,')
add_bullet('vos commentaires éventuels.')

doc.add_heading('Actions disponibles', level=2)
add_table(
    ['Bouton', 'Action'],
    [
        ['Confirmer la commande', 'Envoie définitivement votre commande'],
        ['Modifier le panier', 'Retourne au panier pour modifier les articles'],
    ]
)

doc.add_heading('Après confirmation', level=2)
add_text('Une fois la commande confirmée :')
add_bullet('Un numéro de commande unique vous est attribué (format : CMD-YYYYMMDD-XXXX).')
add_bullet('Un email de confirmation est envoyé à votre adresse.')
add_bullet('Votre panier est automatiquement vidé.')
add_bullet('Vous êtes redirigé vers la page de confirmation.')

add_text(
    'Depuis la page de confirmation, vous pouvez cliquer sur "Retour au catalogue" '
    'pour passer une nouvelle commande.',
)

# ============================================================
# 9. CONSULTER L'HISTORIQUE DES COMMANDES
# ============================================================
doc.add_page_break()
doc.add_heading('9. Consulter l\'historique des commandes', level=1)

doc.add_heading('Accéder à l\'historique', level=2)
add_text(
    'Cliquez sur "Mes commandes" dans la barre de navigation pour accéder '
    'à la liste de vos 50 dernières commandes.'
)

doc.add_heading('Informations affichées', level=2)
add_text('Pour chaque commande, vous voyez :')
add_table(
    ['Colonne', 'Information'],
    [
        ['Numéro', 'Le numéro unique de la commande (ex : CMD-20260212-0001)'],
        ['Date', 'La date et l\'heure de la commande'],
        ['Articles', 'Le nombre d\'articles dans la commande'],
        ['Total HT', 'Le montant total hors taxes'],
        ['Action', 'Bouton "Détail" pour voir les détails'],
    ]
)

doc.add_heading('Détails d\'une commande', level=2)
add_text('En cliquant sur "Détail", vous accédez à la page complète de la commande avec :')
add_bullet('la liste de tous les articles (nom, référence, quantité, prix unitaire, total),')
add_bullet('le total général HT,')
add_bullet('les dates de commande, de départ des camions et de livraison,')
add_bullet('le nom du client et le code client,')
add_bullet('les commentaires associés à la commande.')

add_text('Un bouton "Retour à l\'historique" permet de revenir à la liste.')

# ============================================================
# 10. GERER SON PROFIL
# ============================================================
doc.add_page_break()
doc.add_heading('10. Gérer son profil', level=1)

doc.add_heading('Accéder au profil', level=2)
add_text(
    'Cliquez sur votre nom dans la barre de navigation pour accéder à votre page de profil.'
)

doc.add_heading('Informations affichées', level=2)
add_text('La page de profil affiche vos informations professionnelles :')
add_table(
    ['Information', 'Description'],
    [
        ['Nom client', 'Le nom de votre entreprise'],
        ['Nom d\'utilisateur', 'Votre identifiant de connexion'],
        ['Email', 'Votre adresse email enregistrée'],
        ['Adresse', 'L\'adresse de livraison principale'],
        ['Code postal', 'Le code postal'],
        ['Ville', 'La ville'],
        ['Code client', 'Votre référence dans le système Giffaud'],
    ]
)

add_note(
    'Ces informations proviennent du système central Giffaud. '
    'Pour modifier votre adresse ou votre nom, contactez votre commercial.'
)

doc.add_heading('Modifier son email', level=2)
add_step(1, 'Depuis le profil, cliquez sur "Modifier mon email".')
add_step(2, 'Saisissez votre nouvelle adresse email.')
add_step(3, 'Cliquez sur "Enregistrer".')

doc.add_heading('Modifier son mot de passe', level=2)
add_step(1, 'Depuis le profil, cliquez sur "Modifier mon mot de passe".')
add_step(2, 'Saisissez votre ancien mot de passe.')
add_step(3, 'Saisissez votre nouveau mot de passe.')
add_step(4, 'Confirmez le nouveau mot de passe.')
add_step(5, 'Cliquez sur "Enregistrer".')

add_note(
    'Le mot de passe doit comporter au moins 8 caractères, ne pas être trop courant '
    'et ne pas être entièrement numérique.'
)

# ============================================================
# 11. REINITIALISER SON MOT DE PASSE
# ============================================================
doc.add_page_break()
doc.add_heading('11. Réinitialiser son mot de passe', level=1)

add_text(
    'Si vous avez oublié votre mot de passe, vous pouvez le réinitialiser '
    'depuis la page de connexion.'
)

doc.add_heading('Procédure', level=2)
add_step(1, 'Sur la page de connexion, cliquez sur "Mot de passe oublié ?".')
add_step(2, 'Une fenêtre s\'ouvre. Saisissez votre adresse email.')
add_step(3, 'Cliquez sur "Envoyer le lien".')
add_step(4, 'Consultez votre boîte email (vérifiez aussi les spams).')
add_step(5, 'Cliquez sur le lien reçu par email.')
add_step(6, 'Saisissez votre nouveau mot de passe et confirmez-le.')
add_step(7, 'Retournez sur la page de connexion et connectez-vous avec votre nouveau mot de passe.')

add_note(
    'Le lien de réinitialisation est valable pendant 1 heure. '
    'Passé ce délai, vous devrez refaire une demande.'
)

# ============================================================
# 12. AIDE ET CONTACT
# ============================================================
doc.add_page_break()
doc.add_heading('12. Aide et contact', level=1)

doc.add_heading('En cas de difficulté', level=2)
add_text('Si vous rencontrez un problème avec l\'application :')
add_bullet('Vérifiez que votre navigateur est à jour (Chrome, Firefox, Edge recommandés).')
add_bullet('Essayez de vous déconnecter puis de vous reconnecter.')
add_bullet('Videz le cache de votre navigateur.')
add_bullet('Si le problème persiste, contactez votre commercial Giffaud Groupe.')

doc.add_heading('Informations utiles', level=2)
add_table(
    ['Sujet', 'Que faire'],
    [
        ['Identifiants de connexion', 'Contactez votre commercial'],
        ['Mot de passe oublié', 'Utilisez la fonction "Mot de passe oublié" sur la page de connexion'],
        ['Produits manquants', 'Contactez votre commercial pour vérifier votre catalogue'],
        ['Erreur de prix', 'Contactez votre commercial'],
        ['Problème technique', 'Contactez le support technique Giffaud Groupe'],
        ['Modification d\'adresse', 'Contactez votre commercial'],
    ]
)

doc.add_heading('Mentions légales', level=2)
add_text(
    'Les mentions légales sont accessibles via le lien "Mentions légales" '
    'dans le pied de page de chaque page.'
)

# ============================================================
# RECAPITULATIF RAPIDE
# ============================================================
doc.add_page_break()
doc.add_heading('Récapitulatif rapide', level=1)

add_table(
    ['Action', 'Comment faire'],
    [
        ['Se connecter', 'Saisir identifiant + mot de passe sur la page de connexion'],
        ['Voir le catalogue', 'Menu "Catalogue" ou cliquer sur le logo'],
        ['Chercher un produit', 'Utiliser la barre de recherche dans les filtres'],
        ['Filtrer les produits', 'Cocher des filtres dans le panneau gauche, puis "Filtrer"'],
        ['Ajouter au panier', 'Ajuster la quantité et cliquer sur "Ajouter"'],
        ['Voir le panier', 'Cliquer sur l\'icône panier dans le menu'],
        ['Commander (rapide)', 'Menu "Commander", remplir les quantités, valider'],
        ['Valider une commande', 'Panier → Suivant → Vérifier → Confirmer'],
        ['Voir mes commandes', 'Menu "Mes commandes"'],
        ['Modifier mon profil', 'Cliquer sur son nom dans le menu'],
        ['Changer mon MDP', 'Profil → Modifier mon mot de passe'],
        ['MDP oublié', 'Page de connexion → "Mot de passe oublié ?"'],
        ['Se déconnecter', 'Cliquer sur "Déconnexion" dans le menu'],
    ]
)

# ============================================================
# PIED DE PAGE
# ============================================================
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Document généré le 12/02/2026 — Extranet Giffaud Groupe')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ============================================================
# SAUVEGARDE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), 'GUIDE_UTILISATEUR.docx')
doc.save(output_path)
print(f"Document généré : {output_path}")
